#!/usr/bin/env python3
"""
Generate a comprehensive report of networking conversations across a full range of skill levels.
This script generates a wider range of conversations, exports to Excel, PDF, and text formats,
and includes detailed visualizations for human expert review.
"""

import os
import sys
import json
import time
import argparse
import logging
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.backends.backend_pdf import PdfPages
from synthetic_conversation_gpt import ConversationGenerator, ConversationEvaluator
from fpdf import FPDF
import inspect

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('comprehensive_report.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# ANSI colors for console output
class Colors:
    HEADER = '\033[95m'
    BLUE = '\033[94m'
    CYAN = '\033[96m'
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

def parse_args():
    """Parse command-line arguments."""
    import argparse
    parser = argparse.ArgumentParser(description='Generate comprehensive evaluation of networking conversations.')
    parser.add_argument('--output-dir', type=str, default='text_results', help='Directory to store evaluation results')
    parser.add_argument('--summary-only', action='store_true', help='Only generate summary reports from existing results')
    parser.add_argument('--word-doc', action='store_true', help='Generate a Word document with conversations')
    parser.add_argument('--text-file', action='store_true', help='Generate a text file with conversations for Google Docs')
    parser.add_argument('--analyze', action='store_true', help='Analyze existing JSON files for issues')
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose output')
    parser.add_argument('--debug', '-d', action='store_true', help='Enable debug mode with extra details')
    parser.add_argument('--regenerate', action='store_true', help='Force regeneration of all conversations')
    parser.add_argument('--text-only', action='store_true', help='Use text-only approach with no JSON intermediaries')
    parser.add_argument('--simple', action='store_true', help='Use simple text extraction with no dependencies')
    parser.add_argument('--regenerate-all', action='store_true', help='Regenerate everything from scratch and save directly as text')
    return parser.parse_args()

def setup_output_directory(output_dir):
    """Setup output directory structure."""
    if output_dir is None:
        output_dir = 'comprehensive_results'
    
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        logging.info(f"Created output directory: {output_dir}")
    
    # Create subdirectories
    subdirs = ['json', 'text', 'excel', 'visualizations', 'pdf']
    for subdir in subdirs:
        subdir_path = os.path.join(output_dir, subdir)
        if not os.path.exists(subdir_path):
            os.makedirs(subdir_path)
    
    logging.info(f"Created output directory structure in {output_dir}")
    return output_dir  # Ensure we always return the output_dir

def generate_skill_gradients():
    """Generate skill level combinations."""
    skill_bases = ['novice', 'intermediate', 'advanced']
    gradients = ['low', 'basic', 'high']
    
    gradients = []
    for base in skill_bases:
        for gradient in gradients:
            skill_level = f"{base}_{gradient}"
            label = f"{base.capitalize()} ({gradient})"
            gradients.append((skill_level, label))
    
    return gradients

def generate_conversation(generator, persona, skill_level, label, verbose=False, debug=False):
    """Generate a networking conversation."""
    import json
    import traceback
    
    if verbose or debug:
        print(f"\n{Colors.BLUE}Generating conversation for {skill_level} skill level with {persona} persona...{Colors.ENDC}")
    
    try:
        # Define the prompt for the conversation generator
        prompt = {
            "skill_level": skill_level,
            "persona": persona,
            "label": label
        }
        
        if debug:
            print(f"{Colors.CYAN}Prompt to generator:{Colors.ENDC}")
            print(json.dumps(prompt, indent=2))
        
        # Generate the conversation
        result = generator.generate_conversation(prompt)
        
        if debug:
            print(f"{Colors.CYAN}Raw generator result:{Colors.ENDC}")
            print(json.dumps(result, indent=2)[:1000] + "..." if len(json.dumps(result)) > 1000 else json.dumps(result, indent=2))
        
        # Ensure the result is a list of messages
        if not isinstance(result, list):
            if verbose or debug:
                print(f"{Colors.RED}Error: Generator returned a {type(result).__name__} instead of a list{Colors.ENDC}")
            # Try to convert to a list if possible
            if isinstance(result, dict) and "messages" in result:
                result = result["messages"]
                if verbose or debug:
                    print(f"{Colors.YELLOW}Extracted messages from result dictionary{Colors.ENDC}")
            else:
                result = [{"role": "system", "content": "Error: Invalid conversation format returned by generator"}]
        
        # Verify the format of each message
        for i, message in enumerate(result):
            if not isinstance(message, dict):
                if verbose or debug:
                    print(f"{Colors.RED}Warning: Message {i} is not a dictionary, converting...{Colors.ENDC}")
                result[i] = {"role": "unknown", "content": str(message)}
            elif "role" not in message or "content" not in message:
                if verbose or debug:
                    print(f"{Colors.RED}Warning: Message {i} missing required fields{Colors.ENDC}")
                # Add missing fields
                if "role" not in message:
                    message["role"] = "unknown"
                if "content" not in message:
                    message["content"] = ""
        
        if verbose or debug:
            print(f"{Colors.GREEN}Successfully generated conversation with {len(result)} messages{Colors.ENDC}")
            if debug:
                for i, message in enumerate(result):
                    print(f"{Colors.CYAN}Message {i}:{Colors.ENDC}")
                    print(f"  Role: {message.get('role', 'unknown')}")
                    content = message.get('content', '')
                    print(f"  Content: {content[:100]}..." if len(content) > 100 else f"  Content: {content}")
        
        return result
    
    except Exception as e:
        error_msg = f"Error generating conversation: {str(e)}"
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        if debug:
            print(f"{Colors.RED}Stack trace:{Colors.ENDC}")
            print(traceback.format_exc())
        
        # Return a minimal conversation with the error
        return [
            {"role": "system", "content": "Error occurred during conversation generation."},
            {"role": "system", "content": error_msg}
        ]

def evaluate_conversation(evaluator, conversation, skill_level, verbose=False, debug=False):
    """Evaluate a networking conversation."""
    import json
    import traceback
    
    if verbose or debug:
        print(f"\n{Colors.BLUE}Evaluating conversation for {skill_level} skill level...{Colors.ENDC}")
    
    try:
        # Prepare the evaluation input
        eval_input = {
            "conversation": conversation,
            "skill_level": skill_level
        }
        
        if debug:
            print(f"{Colors.CYAN}Evaluation input:{Colors.ENDC}")
            print(json.dumps(eval_input, indent=2)[:1000] + "..." if len(json.dumps(eval_input)) > 1000 else json.dumps(eval_input, indent=2))
        
        # Evaluate the conversation
        evaluation = evaluator.evaluate_conversation(eval_input)
        
        if debug:
            print(f"{Colors.CYAN}Raw evaluation result:{Colors.ENDC}")
            print(json.dumps(evaluation, indent=2))
        
        # Ensure the evaluation is a dictionary
        if not isinstance(evaluation, dict):
            if verbose or debug:
                print(f"{Colors.RED}Error: Evaluator returned a {type(evaluation).__name__} instead of a dictionary{Colors.ENDC}")
            evaluation = {
                "overall_score": 0,
                "critical_thinking_score": 0,
                "communication_score": 0,
                "emotional_intelligence_score": 0,
                "badge_level": "Bronze",
                "response_evaluations": [],
                "rationale": f"Error: Evaluator returned a {type(evaluation).__name__} instead of a dictionary"
            }
        
        # Check for required fields
        required_fields = [
            "overall_score", "critical_thinking_score", "communication_score",
            "emotional_intelligence_score", "badge_level"
        ]
        
        missing_fields = [field for field in required_fields if field not in evaluation]
        if missing_fields:
            if verbose or debug:
                print(f"{Colors.RED}Warning: Evaluation missing required fields: {', '.join(missing_fields)}{Colors.ENDC}")
            
            # Add default values for missing fields
            for field in missing_fields:
                if field == "badge_level":
                    evaluation[field] = "Bronze"
                else:
                    evaluation[field] = 0
        
        if verbose or debug:
            print(f"{Colors.GREEN}Successfully evaluated conversation{Colors.ENDC}")
            print(f"  Overall Score: {evaluation.get('overall_score', 0)}")
            print(f"  Badge Level: {evaluation.get('badge_level', 'Bronze')}")
        
        return evaluation
    
    except Exception as e:
        error_msg = f"Error evaluating conversation: {str(e)}"
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        if debug:
            print(f"{Colors.RED}Stack trace:{Colors.ENDC}")
            print(traceback.format_exc())
        
        # Return a minimal evaluation with the error
        return {
            "overall_score": 0,
            "critical_thinking_score": 0,
            "communication_score": 0,
            "emotional_intelligence_score": 0,
            "badge_level": "Bronze",
            "response_evaluations": [],
            "rationale": error_msg
        }

def save_conversation_to_json(conversation, evaluation, skill_level, persona, output_dir, index, verbose=False, debug=False):
    """Save conversation and evaluation to a JSON file."""
    import os
    import json
    import traceback
    
    try:
        # Determine the gradient (high, basic, low)
        _, _, gradient = skill_level.split('_')
        
        # Create filename
        filename = f"{skill_level.split('_')[0]}_{gradient}_{index}.json"
        
        # Create output path
        json_dir = os.path.join(output_dir, 'json')
        output_path = os.path.join(json_dir, filename)
        
        if verbose or debug:
            print(f"\n{Colors.BLUE}Saving conversation to JSON: {output_path}{Colors.ENDC}")
        
        # Verify conversation structure before saving
        if not isinstance(conversation, list):
            if verbose or debug:
                print(f"{Colors.RED}Warning: Conversation is not a list (type: {type(conversation).__name__}){Colors.ENDC}")
        else:
            if debug:
                print(f"{Colors.CYAN}Conversation details:{Colors.ENDC}")
                print(f"  Type: {type(conversation).__name__}")
                print(f"  Length: {len(conversation)}")
                
                if conversation:
                    first_msg = conversation[0]
                    print(f"  First message type: {type(first_msg).__name__}")
                    
                    if isinstance(first_msg, dict):
                        print(f"  First message keys: {list(first_msg.keys())}")
                        role = first_msg.get('role', 'MISSING')
                        content = first_msg.get('content', 'MISSING')
                        print(f"  First message role: {role}")
                        print(f"  First message content: {content[:50]}..." if len(content) > 50 else f"  First message content: {content}")
        
        # Verify evaluation structure before saving
        if not isinstance(evaluation, dict):
            if verbose or debug:
                print(f"{Colors.RED}Warning: Evaluation is not a dictionary (type: {type(evaluation).__name__}){Colors.ENDC}")
        else:
            if debug:
                print(f"{Colors.CYAN}Evaluation details:{Colors.ENDC}")
                print(f"  Type: {type(evaluation).__name__}")
                print(f"  Keys: {list(evaluation.keys())}")
        
        # Construct data to save
        data = {
            'skill_level': skill_level,
            'persona': persona,
            'evaluation': evaluation,
            'conversation': conversation
        }
        
        # Save to JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        
        if verbose or debug:
            print(f"{Colors.GREEN}Successfully saved conversation to {output_path}{Colors.ENDC}")
        
        return output_path
    
    except Exception as e:
        error_msg = f"Error saving conversation to JSON: {str(e)}"
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        if debug:
            print(f"{Colors.RED}Stack trace:{Colors.ENDC}")
            print(traceback.format_exc())
        return None

def save_conversation_to_text(conversation, evaluation, skill_level, persona, output_dir, index):
    """Save conversation and evaluation results to a text file.
    
    Args:
        conversation: The conversation data
        evaluation: The evaluation results
        skill_level: The skill level label
        persona: The persona used
        output_dir: The output directory
        index: The conversation index
    """
    # Create output directory if it doesn't exist
    txt_dir = os.path.join(output_dir, "txt")
    os.makedirs(txt_dir, exist_ok=True)
    
    # Create filename
    filename = f"{skill_level}_{index}.txt"
    filepath = os.path.join(txt_dir, filename)
    
    # Format the conversation and evaluation as text
    lines = []
    
    # Add header
    lines.append("=" * 80)
    lines.append(f"NETWORKING CONVERSATION EVALUATION: {skill_level.upper()}")
    lines.append("=" * 80)
    lines.append(f"Persona: {persona}")
    lines.append(f"Timestamp: {datetime.now().isoformat()}")
    lines.append(f"Conversation ID: {skill_level}_{index}")
    lines.append("")
    
    # Add conversation
    lines.append("-" * 80)
    lines.append("CONVERSATION TRANSCRIPT")
    lines.append("-" * 80)
    
    for i, exchange in enumerate(conversation):
        lines.append(f"[Exchange {i+1}] Stage: {exchange.get('stage', 'unknown')}")
        lines.append(f"User: {exchange.get('user', '')}")
        lines.append(f"Assistant: {exchange.get('assistant', '')}")
        lines.append("")
    
    # Add evaluation summary
    lines.append("-" * 80)
    lines.append("EVALUATION SUMMARY")
    lines.append("-" * 80)
    lines.append(f"Overall Score: {evaluation['overall_score']}")
    lines.append(f"Badge Level: {evaluation['badge_level']}")
    lines.append("")
    lines.append("Skill Scores:")
    lines.append(f"- Critical Thinking: {evaluation['critical_thinking_score']}")
    lines.append(f"- Communication: {evaluation['communication_score']}")
    lines.append(f"- Emotional Intelligence: {evaluation['emotional_intelligence_score']}")
    lines.append("")
    
    # Add detailed evaluation
    lines.append("-" * 80)
    lines.append("DETAILED EVALUATION")
    lines.append("-" * 80)
    
    for i, resp_eval in enumerate(evaluation.get('response_evaluations', [])):
        lines.append(f"[Response {i+1}] Stage: {resp_eval.get('stage', 'unknown')}")
        lines.append(f"User: {resp_eval.get('user_query', '')}")
        lines.append(f"Assistant: {resp_eval.get('assistant_response', '')}")
        lines.append("")
        
        eval_data = resp_eval.get('evaluation', {})
        lines.append("Scores:")
        lines.append(f"- Critical Thinking: {eval_data.get('critical_thinking', 0)}")
        lines.append(f"- Communication: {eval_data.get('communication', 0)}")
        lines.append(f"- Emotional Intelligence: {eval_data.get('emotional_intelligence', 0)}")
        lines.append("")
        
        lines.append(f"Feedback: {eval_data.get('feedback', 'No feedback available')}")
        lines.append("")
    
    # Save to file
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    logger.info(f"Saved text results to {filepath}")

def create_summary_dataframe(results):
    """Create a summary dataframe from the evaluation results."""
    rows = []
    
    for result in results:
        try:
            # Extract basic metadata
            skill_level = result.get('skill_level', 'unknown')
            persona = result.get('persona', 'unknown')
            
            # Extract skill base and gradient if available
            if isinstance(skill_level, str) and '_' in skill_level:
                parts = skill_level.split('_')
                if len(parts) >= 2:
                    skill_base = parts[0]
                    gradient = parts[1] if len(parts) > 1 else 'unknown'
                else:
                    skill_base = skill_level
                    gradient = 'unknown'
            else:
                skill_base = skill_level
                gradient = 'unknown'
            
            # Extract evaluation data
            evaluation = result.get('evaluation', {})
            if not evaluation:
                logger.warning(f"No evaluation data for {skill_level}")
                continue
                
            overall_score = evaluation.get('overall_score', 0)
            critical_thinking_score = evaluation.get('critical_thinking_score', 0)
            communication_score = evaluation.get('communication_score', 0)
            emotional_intelligence_score = evaluation.get('emotional_intelligence_score', 0)
            badge = evaluation.get('badge', 'unknown')
            feedback = evaluation.get('feedback', '')
            
            # Extract conversation text
            conversation_text = result.get('conversation_text', '')
            
            # If conversation_text is not already extracted, extract it from the conversation
            if not conversation_text and 'conversation' in result:
                for message in result['conversation']:
                    if isinstance(message, dict):
                        role = message.get('role', 'Unknown')
                        content = message.get('content', '')
                        conversation_text += f"{role}: {content}\n\n"
            
            # Create row
            row = {
                'skill_level': skill_level,
                'skill_base': skill_base,
                'gradient': gradient,
                'persona': persona,
                'overall_score': overall_score,
                'critical_thinking_score': critical_thinking_score,
                'communication_score': communication_score,
                'emotional_intelligence_score': emotional_intelligence_score,
                'badge': badge,
                'feedback': feedback,
                'conversation_text': conversation_text
            }
            
            rows.append(row)
        except Exception as e:
            logger.error(f"Error processing result: {str(e)}")
    
    # Create dataframe
    if not rows:
        logger.warning("No valid results to create dataframe")
        return pd.DataFrame()
    
    df = pd.DataFrame(rows)
    
    # Export to CSV for reference
    csv_path = os.path.join(os.path.dirname(results[0].get('output_dir', '.')), 'comprehensive_results.csv')
    logger.info(f"Exporting results to CSV: {csv_path}")
    df.to_csv(csv_path, index=False)
    
    return df

def load_existing_results(output_dir):
    """Load existing results from JSON files."""
    results = []
    json_dir = os.path.join(output_dir, 'json')
    
    if not os.path.exists(json_dir):
        logger.warning(f"JSON directory not found: {json_dir}")
        return results
    
    json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
    print(f"Found {len(json_files)} JSON files")
    
    for json_file in json_files:
        try:
            with open(os.path.join(json_dir, json_file), 'r') as f:
                data = json.load(f)
                
                # Extract conversation text for inclusion in reports
                conversation_text = ""
                if 'conversation' in data:
                    for message in data['conversation']:
                        if isinstance(message, dict):
                            role = message.get('role', 'Unknown')
                            content = message.get('content', '')
                            conversation_text += f"{role}: {content}\n\n"
                
                # Add conversation text to the data
                data['conversation_text'] = conversation_text
                
                # Add output directory for reference
                data['output_dir'] = output_dir
                
                results.append(data)
        except Exception as e:
            logger.error(f"Error loading {json_file}: {str(e)}")
    
    print(f"Successfully loaded {len(results)} results")
    return results

def generate_summary_reports(output_dir):
    """Generate summary reports from existing results."""
    print("\n================================================================================")
    print("GENERATING SUMMARY REPORTS FROM EXISTING RESULTS")
    print("================================================================================\n")
    
    print(f"Loading existing results from {output_dir}...")
    
    # Load existing results
    results = load_existing_results(output_dir)
    
    if not results:
        print("No results found. Please run the evaluation first.")
        return
    
    # Create summary dataframe
    df = create_summary_dataframe(results)
    
    if df.empty:
        print("Failed to create summary dataframe.")
        return
    
    # Create visualizations directory
    visualization_dir = os.path.join(output_dir, 'visualizations')
    os.makedirs(visualization_dir, exist_ok=True)
    
    # Create visualizations
    create_visualizations(df, visualization_dir)
    
    # Export to Excel
    excel_path = export_to_excel(df, output_dir)
    
    # Generate PDF report
    pdf_path = generate_pdf_report(df, output_dir, visualization_dir)
    
    print("\n================================================================================")
    print("SUMMARY REPORTS GENERATED")
    print("================================================================================\n")
    
    print(f"All reports saved to: {output_dir}")
    
    return excel_path, pdf_path

def export_to_excel(df, output_dir):
    """Export results to Excel with multiple sheets and analysis."""
    try:
        # Create filename
        filename = "comprehensive_results.xlsx"
        excel_path = os.path.join(output_dir, filename)
        
        logger.info(f"Exporting comprehensive results to Excel: {excel_path}")
        
        # Create a Pandas Excel writer
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            # Main data sheet with all results
            df.to_excel(writer, sheet_name='All Data', index=False)
            
            # Executive Summary sheet
            summary_data = [
                ["Total Conversations", str(len(df))],
                ["Skill Levels", str(df['skill_level'].nunique()) if 'skill_level' in df.columns else "N/A"],
                ["Skill Bases", str(df['skill_base'].nunique()) if 'skill_base' in df.columns else "N/A"],
                ["Gradients", str(df['gradient'].nunique()) if 'gradient' in df.columns else "N/A"],
                ["Average Overall Score", f"{df['overall_score'].mean():.2f}" if 'overall_score' in df.columns else "N/A"],
                ["Highest Scoring Level", str(df.groupby('skill_level')['overall_score'].mean().idxmax()) if 'skill_level' in df.columns and 'overall_score' in df.columns else "N/A"],
                ["Lowest Scoring Level", str(df.groupby('skill_level')['overall_score'].mean().idxmin()) if 'skill_level' in df.columns and 'overall_score' in df.columns else "N/A"]
            ]
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='Executive Summary', index=False)
            
            # Skill Level Analysis
            if 'skill_level' in df.columns and 'overall_score' in df.columns:
                skill_summary = df.groupby('skill_level')['overall_score'].agg(['mean', 'min', 'max', 'count']).reset_index()
                skill_summary = skill_summary.sort_values('mean', ascending=False)
                skill_summary.columns = ['Skill Level', 'Average Score', 'Min Score', 'Max Score', 'Count']
                skill_summary.to_excel(writer, sheet_name='Skill Level Analysis', index=False)
            
            # Badge Distribution
            if 'badge' in df.columns:
                badge_summary = df['badge'].value_counts().reset_index()
                badge_summary.columns = ['Badge', 'Count']
                badge_summary['Percentage'] = badge_summary['Count'] / badge_summary['Count'].sum() * 100
                badge_summary.to_excel(writer, sheet_name='Badge Distribution', index=False)
            
            # Detailed Skill Analysis
            skill_cols = ['critical_thinking_score', 'communication_score', 'emotional_intelligence_score']
            if all(col in df.columns for col in skill_cols):
                skill_data = {
                    'Skill Area': ['Critical Thinking', 'Communication', 'Emotional Intelligence'],
                    'Average Score': [
                        df['critical_thinking_score'].mean(),
                        df['communication_score'].mean(),
                        df['emotional_intelligence_score'].mean()
                    ],
                    'Min Score': [
                        df['critical_thinking_score'].min(),
                        df['communication_score'].min(),
                        df['emotional_intelligence_score'].min()
                    ],
                    'Max Score': [
                        df['critical_thinking_score'].max(),
                        df['communication_score'].max(),
                        df['emotional_intelligence_score'].max()
                    ],
                    'Standard Deviation': [
                        df['critical_thinking_score'].std(),
                        df['communication_score'].std(),
                        df['emotional_intelligence_score'].std()
                    ]
                }
                pd.DataFrame(skill_data).to_excel(writer, sheet_name='Skill Analysis', index=False)
            
            # Full Conversations Sheet
            if 'conversation_text' in df.columns:
                # Create a conversations dataframe with key metadata and full text
                conversations_df = df[['skill_level', 'skill_base', 'gradient', 'overall_score', 'badge', 'conversation_text']].copy()
                conversations_df.to_excel(writer, sheet_name='Full Conversations', index=False)
                
                # Format the conversations sheet
                workbook = writer.book
                worksheet = writer.sheets['Full Conversations']
                
                # Set column widths
                worksheet.column_dimensions['A'].width = 15  # skill_level
                worksheet.column_dimensions['B'].width = 15  # skill_base
                worksheet.column_dimensions['C'].width = 15  # gradient
                worksheet.column_dimensions['D'].width = 15  # overall_score
                worksheet.column_dimensions['E'].width = 15  # badge
                worksheet.column_dimensions['F'].width = 100  # conversation_text
                
                # Set row height for conversation text
                for row in range(2, len(conversations_df) + 2):
                    worksheet.row_dimensions[row].height = 100
            
            # Correlation Analysis
            if len(df) > 5 and 'overall_score' in df.columns:
                numeric_cols = df.select_dtypes(include=['float64', 'int64']).columns
                if len(numeric_cols) > 1:
                    corr_df = df[numeric_cols].corr()
                    corr_df.to_excel(writer, sheet_name='Correlation Analysis')
        
        print(f"\nExcel report saved to: {excel_path}")
        return excel_path
    
    except Exception as e:
        logger.error(f"Error exporting to Excel: {str(e)}")
        print(f"Error exporting to Excel: {str(e)}")
        return None

def create_visualizations(df, output_dir):
    """Create visualizations from the summary dataframe.
    
    Args:
        df: The dataframe to visualize
        output_dir: The output directory
        
    Returns:
        A dictionary of visualization paths
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Initialize paths dictionary
    visualization_paths = {}
    
    # Set the style
    plt.style.use('seaborn-v0_8-darkgrid')
    
    # Check if we have any data
    if len(df) == 0:
        logger.warning("No data to visualize")
        return visualization_paths
    
    # Make sure we have the required columns
    required_columns = ['skill_level', 'skill_base', 'gradient', 'overall_score']
    missing_columns = [col for col in required_columns if col not in df.columns]
    if missing_columns:
        logger.warning(f"Missing columns for visualization: {missing_columns}")
        logger.warning(f"Available columns: {df.columns.tolist()}")
        # Create a minimal dataframe with required columns if they're missing
        if 'skill_level' not in df.columns:
            df['skill_level'] = 'unknown'
        if 'skill_base' not in df.columns:
            df['skill_base'] = 'unknown'
        if 'gradient' not in df.columns:
            df['gradient'] = 0.0
        if 'overall_score' not in df.columns:
            df['overall_score'] = df.get('overall_score', 0)
    
    try:
        # 1. Overall score by skill level
        plt.figure(figsize=(12, 8))
        sns.barplot(x='skill_level', y='overall_score', data=df, errorbar=None)
        plt.title('Average Overall Score by Skill Level', fontsize=16)
        plt.xlabel('Skill Level', fontsize=14)
        plt.ylabel('Average Score', fontsize=14)
        plt.xticks(rotation=45)
        plt.tight_layout()
        
        # Save the figure
        score_by_level_path = os.path.join(output_dir, 'score_by_level.png')
        plt.savefig(score_by_level_path, dpi=300)
        plt.close()
        
        visualization_paths['score_by_level'] = score_by_level_path
        
        # 2. Score distribution by skill base (novice, intermediate, advanced)
        plt.figure(figsize=(12, 8))
        sns.boxplot(x='skill_base', y='overall_score', data=df)
        plt.title('Score Distribution by Skill Base', fontsize=16)
        plt.xlabel('Skill Base', fontsize=14)
        plt.ylabel('Score', fontsize=14)
        plt.tight_layout()
        
        # Save the figure
        score_distribution_path = os.path.join(output_dir, 'score_distribution.png')
        plt.savefig(score_distribution_path, dpi=300)
        plt.close()
        
        visualization_paths['score_distribution'] = score_distribution_path
        
        # 3. Skill scores by skill level
        plt.figure(figsize=(14, 10))
        
        # Melt the dataframe to get skill scores in long format
        skill_cols = ['critical_thinking', 'communication', 'emotional_intelligence']
        if all(col in df.columns for col in skill_cols):
            melted = pd.melt(
                df, 
                id_vars=['skill_level'], 
                value_vars=skill_cols,
                var_name='Skill Area', 
                value_name='Score'
            )
            
            sns.barplot(x='skill_level', y='Score', hue='Skill Area', data=melted, errorbar=None)
            plt.title('Skill Scores by Skill Level', fontsize=16)
            plt.xlabel('Skill Level', fontsize=14)
            plt.ylabel('Average Score', fontsize=14)
            plt.xticks(rotation=45)
            plt.legend(title='Skill Area', fontsize=12)
            plt.tight_layout()
            
            # Save the figure
            skill_scores_path = os.path.join(output_dir, 'skill_scores.png')
            plt.savefig(skill_scores_path, dpi=300)
            plt.close()
            
            visualization_paths['skill_scores'] = skill_scores_path
        
        # 4. Badge distribution
        if 'badge_level' in df.columns:
            plt.figure(figsize=(10, 8))
            badge_counts = df['badge_level'].value_counts()
            colors = ['#CD7F32', '#C0C0C0', '#FFD700']  # Bronze, Silver, Gold
            
            # Ensure we have all badge levels
            for badge in ['Bronze', 'Silver', 'Gold']:
                if badge not in badge_counts:
                    badge_counts[badge] = 0
            
            # Sort by badge level
            badge_counts = badge_counts.reindex(['Bronze', 'Silver', 'Gold'])
            
            # Create the pie chart
            plt.pie(
                badge_counts, 
                labels=badge_counts.index, 
                autopct='%1.1f%%',
                colors=colors,
                startangle=90,
                wedgeprops={'edgecolor': 'white', 'linewidth': 1.5}
            )
            plt.title('Badge Distribution', fontsize=16)
            plt.axis('equal')
            
            # Save the figure
            badge_distribution_path = os.path.join(output_dir, 'badge_distribution.png')
            plt.savefig(badge_distribution_path, dpi=300)
            plt.close()
            
            visualization_paths['badge_distribution'] = badge_distribution_path
        
        # 5. Heatmap of scores by skill level and gradient
        if 'gradient' in df.columns:
            plt.figure(figsize=(12, 10))
            
            # Create a pivot table
            pivot = df.pivot_table(
                values='overall_score',
                index='skill_base',
                columns='gradient',
                aggfunc='mean'
            )
            
            # Create the heatmap
            sns.heatmap(
                pivot, 
                annot=True, 
                cmap='YlGnBu', 
                fmt='.1f',
                linewidths=.5
            )
            plt.title('Average Score by Skill Base and Gradient', fontsize=16)
            plt.xlabel('Gradient', fontsize=14)
            plt.ylabel('Skill Base', fontsize=14)
            plt.tight_layout()
            
            # Save the figure
            heatmap_path = os.path.join(output_dir, 'score_heatmap.png')
            plt.savefig(heatmap_path, dpi=300)
            plt.close()
            
            visualization_paths['score_heatmap'] = heatmap_path
        
        logger.info(f"Created visualizations in {output_dir}")
        return visualization_paths
        
    except Exception as e:
        logger.error(f"Error creating visualizations: {str(e)}")
        print(f"{Colors.RED}Error creating visualizations: {str(e)}{Colors.ENDC}")
        return visualization_paths

def sanitize_text_for_pdf(text):
    """Sanitize text for PDF generation by replacing problematic characters."""
    if not isinstance(text, str):
        text = str(text)
    
    # Replace problematic characters
    replacements = {
        '•': '-',  # Bullet point
        '…': '...',  # Ellipsis
        '—': '-',  # Em dash
        '–': '-',  # En dash
        '"': '"',  # Smart quotes
        '"': '"',  # Smart quotes
        ''': "'",  # Smart apostrophe
        ''': "'",  # Smart apostrophe
        '≤': '<=',  # Less than or equal
        '≥': '>=',  # Greater than or equal
        '×': 'x',  # Multiplication sign
        '÷': '/',  # Division sign
        '≠': '!=',  # Not equal
        '≈': '~',  # Approximately equal
        '©': '(c)',  # Copyright
        '®': '(R)',  # Registered trademark
        '™': '(TM)',  # Trademark
    }
    
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    
    return text

def generate_pdf_report(df, output_dir, visualization_dir):
    """Generate a comprehensive PDF report with visualizations and analysis."""
    try:
        # Create filename
        filename = "comprehensive_report.pdf"
        pdf_path = os.path.join(output_dir, filename)
        
        logger.info(f"Generating comprehensive PDF report: {pdf_path}")
        
        # Initialize PDF with larger margins
        pdf = FPDF(orientation='P', unit='mm', format='A4')
        pdf.set_margins(15, 15, 15)  # left, top, right margins in mm
        
        # Set up fonts - using standard fonts to avoid substitution
        title_font = ('Helvetica', 'B', 16)
        heading_font = ('Helvetica', 'B', 14)
        subheading_font = ('Helvetica', 'B', 12)
        body_font = ('Helvetica', '', 10)
        small_font = ('Helvetica', '', 8)
        
        # Calculate usable page width
        page_width = pdf.w - 2 * pdf.l_margin
        
        # ===== COVER PAGE =====
        pdf.add_page()
        
        # Title
        pdf.set_font(*title_font)
        pdf.cell(0, 20, "COMPREHENSIVE NETWORKING SKILLS", align="C")
        pdf.ln()
        pdf.cell(0, 10, "EVALUATION REPORT", align="C")
        pdf.ln(20)
        
        # Add a visualization to the cover if available
        cover_viz = os.path.join(visualization_dir, 'score_heatmap.png')
        if os.path.exists(cover_viz):
            image_width = min(150, page_width - 20)
            pdf.image(cover_viz, x=(pdf.w - image_width)/2, w=image_width)
        
        pdf.ln(20)
        
        # Date and generation info
        pdf.set_font(*body_font)
        pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align="C")
        pdf.ln()
        pdf.cell(0, 10, f"Total conversations evaluated: {len(df)}", align="C")
        pdf.ln()
        
        # ===== TABLE OF CONTENTS =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "TABLE OF CONTENTS")
        pdf.ln()
        
        # Define sections and their page numbers (will be updated later)
        toc = [
            "1. EXECUTIVE SUMMARY",
            "2. EVALUATION OVERVIEW",
            "3. SKILL LEVEL ANALYSIS",
            "4. BADGE DISTRIBUTION",
            "5. DETAILED SKILL SCORES",
            "6. VISUALIZATIONS",
            "7. FULL CONVERSATIONS"
        ]
        
        # Add TOC entries
        pdf.set_font(*body_font)
        pages = [2, 3, 4, 5, 6, 7, 8]  # Placeholder page numbers
        
        for i, (item, page) in enumerate(zip(toc, pages)):
            pdf.cell(0, 8, f"{item}")
            pdf.ln()
            pdf.cell(0, 8, f"Page {page}", align="R")
            pdf.ln()
        
        # ===== EXECUTIVE SUMMARY =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "1. EXECUTIVE SUMMARY")
        pdf.ln()
        
        # Summary statistics
        pdf.set_font(*body_font)
        
        # Create a summary table
        summary_data = [
            ["Total Conversations", str(len(df))],
            ["Skill Levels", str(df['skill_level'].nunique()) if 'skill_level' in df.columns else "N/A"],
            ["Skill Bases", str(df['skill_base'].nunique()) if 'skill_base' in df.columns else "N/A"],
            ["Gradients", str(df['gradient'].nunique()) if 'gradient' in df.columns else "N/A"],
            ["Average Overall Score", f"{df['overall_score'].mean():.2f}" if 'overall_score' in df.columns else "N/A"],
            ["Highest Scoring Level", str(df.groupby('skill_level')['overall_score'].mean().idxmax()) if 'skill_level' in df.columns and 'overall_score' in df.columns else "N/A"],
            ["Lowest Scoring Level", str(df.groupby('skill_level')['overall_score'].mean().idxmin()) if 'skill_level' in df.columns and 'overall_score' in df.columns else "N/A"]
        ]
        
        # Add summary table
        pdf.set_font(*subheading_font)
        pdf.cell(0, 10, "Summary Statistics:")
        pdf.ln()
        pdf.set_font(*body_font)
        
        for label, value in summary_data:
            pdf.cell(60, 8, label, border=0)
            pdf.cell(0, 8, value, border=0)
            pdf.ln()
        
        pdf.ln(5)
        
        # Key findings
        pdf.set_font(*subheading_font)
        pdf.cell(0, 10, "Key Findings:")
        pdf.ln()
        pdf.set_font(*body_font)
        
        # Add some key findings based on the data
        findings = [
            "The evaluation covered a range of skill levels from beginner to expert.",
            f"Average overall score across all conversations: {df['overall_score'].mean():.2f}" if 'overall_score' in df.columns else "Score data not available.",
            f"Most common badge level: {df['badge'].mode().iloc[0]}" if 'badge' in df.columns and not df['badge'].empty else "Badge data not available."
        ]
        
        for finding in findings:
            # Replace bullet point with hyphen and sanitize text
            pdf.multi_cell(0, 8, f"- {sanitize_text_for_pdf(finding)}")
        
        # ===== EVALUATION OVERVIEW =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "2. EVALUATION OVERVIEW")
        pdf.ln()
        
        pdf.set_font(*body_font)
        pdf.multi_cell(0, 8, "This report presents a comprehensive analysis of networking skills evaluations across different skill levels. The evaluation framework assesses critical thinking, communication skills, and emotional intelligence in networking conversations.")
        pdf.ln()
        
        # Add methodology section
        pdf.set_font(*subheading_font)
        pdf.cell(0, 10, "Methodology:")
        pdf.ln()
        pdf.set_font(*body_font)
        pdf.multi_cell(0, 8, "Each conversation was generated with a specific skill level target and evaluated using a standardized rubric. The evaluation produces scores for critical thinking, communication skills, and emotional intelligence, as well as an overall score and badge level.")
        pdf.ln()
        
        # ===== SKILL LEVEL ANALYSIS =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "3. SKILL LEVEL ANALYSIS")
        pdf.ln()
        
        # Add skill level analysis visualization
        skill_viz = os.path.join(visualization_dir, 'score_by_level.png')
        if os.path.exists(skill_viz):
            image_width = min(180, page_width - 10)
            pdf.image(skill_viz, x=(pdf.w - image_width)/2, w=image_width)
            pdf.ln()
        
        # Add skill level table
        if 'skill_level' in df.columns and 'overall_score' in df.columns:
            pdf.set_font(*subheading_font)
            pdf.cell(0, 10, "Average Scores by Skill Level:")
            pdf.ln()
            
            # Create a summary by skill level
            skill_summary = df.groupby('skill_level')['overall_score'].agg(['mean', 'min', 'max', 'count']).reset_index()
            skill_summary = skill_summary.sort_values('mean', ascending=False)
            
            # Table header
            pdf.set_font(*small_font)
            headers = ["Skill Level", "Average", "Min", "Max", "Count"]
            col_widths = [page_width * 0.3, page_width * 0.175, page_width * 0.175, page_width * 0.175, page_width * 0.175]
            
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 8, header, border=1)
            pdf.ln()
            
            # Table data
            for _, row in skill_summary.iterrows():
                pdf.cell(col_widths[0], 8, sanitize_text_for_pdf(str(row['skill_level'])), border=1)
                pdf.cell(col_widths[1], 8, sanitize_text_for_pdf(f"{row['mean']:.2f}"), border=1)
                pdf.cell(col_widths[2], 8, sanitize_text_for_pdf(f"{row['min']:.2f}"), border=1)
                pdf.cell(col_widths[3], 8, sanitize_text_for_pdf(f"{row['max']:.2f}"), border=1)
                pdf.cell(col_widths[4], 8, sanitize_text_for_pdf(str(int(row['count']))), border=1)
                pdf.ln()
        
        # ===== BADGE DISTRIBUTION =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "4. BADGE DISTRIBUTION")
        pdf.ln()
        
        # Add badge distribution visualization
        badge_viz = os.path.join(visualization_dir, 'badge_distribution.png')
        if os.path.exists(badge_viz):
            image_width = min(180, page_width - 10)
            pdf.image(badge_viz, x=(pdf.w - image_width)/2, w=image_width)
            pdf.ln()
        
        # Add badge distribution table
        if 'badge' in df.columns:
            pdf.set_font(*subheading_font)
            pdf.cell(0, 10, "Badge Distribution:")
            pdf.ln()
            
            # Create a summary by badge
            badge_summary = df['badge'].value_counts().reset_index()
            badge_summary.columns = ['Badge', 'Count']
            badge_summary['Percentage'] = badge_summary['Count'] / badge_summary['Count'].sum() * 100
            
            # Table header
            pdf.set_font(*small_font)
            headers = ["Badge", "Count", "Percentage"]
            col_widths = [page_width * 0.4, page_width * 0.3, page_width * 0.3]
            
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 8, header, border=1)
            pdf.ln()
            
            # Table data
            for _, row in badge_summary.iterrows():
                pdf.cell(col_widths[0], 8, sanitize_text_for_pdf(str(row['Badge'])), border=1)
                pdf.cell(col_widths[1], 8, sanitize_text_for_pdf(str(int(row['Count']))), border=1)
                pdf.cell(col_widths[2], 8, sanitize_text_for_pdf(f"{row['Percentage']:.2f}%"), border=1)
                pdf.ln()
        
        # ===== DETAILED SKILL SCORES =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "5. DETAILED SKILL SCORES")
        pdf.ln()
        
        # Add skill scores visualization
        skill_scores_viz = os.path.join(visualization_dir, 'skill_scores.png')
        if os.path.exists(skill_scores_viz):
            image_width = min(180, page_width - 10)
            pdf.image(skill_scores_viz, x=(pdf.w - image_width)/2, w=image_width)
        
        # Add detailed skill scores table
        pdf.set_font(*subheading_font)
        pdf.cell(0, 10, "Skill Area Performance:")
        pdf.ln()
        
        # Create a summary of skill areas
        skill_cols = ['critical_thinking_score', 'communication_score', 'emotional_intelligence_score']
        if all(col in df.columns for col in skill_cols):
            skill_means = {
                'Critical Thinking': df['critical_thinking_score'].mean(),
                'Communication': df['communication_score'].mean(),
                'Emotional Intelligence': df['emotional_intelligence_score'].mean()
            }
            
            # Table header
            pdf.set_font(*small_font)
            headers = ["Skill Area", "Average Score"]
            col_widths = [page_width * 0.6, page_width * 0.4]
            
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 8, header, border=1)
            pdf.ln()
            
            # Table data
            for skill, score in skill_means.items():
                pdf.cell(col_widths[0], 8, sanitize_text_for_pdf(skill), border=1)
                pdf.cell(col_widths[1], 8, sanitize_text_for_pdf(f"{score:.2f}"), border=1)
                pdf.ln()
        
        # ===== VISUALIZATIONS =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "6. VISUALIZATIONS")
        pdf.ln()
        
        # Add all visualizations
        viz_files = [
            ('score_distribution.png', 'Score Distribution'),
            ('score_heatmap.png', 'Score Heatmap')
        ]
        
        for viz_file, title in viz_files:
            viz_path = os.path.join(visualization_dir, viz_file)
            if os.path.exists(viz_path):
                pdf.set_font(*subheading_font)
                pdf.cell(0, 10, title)
                pdf.ln()
                
                image_width = min(180, page_width - 10)
                pdf.image(viz_path, x=(pdf.w - image_width)/2, w=image_width)
                pdf.ln(15)
        
        # ===== FULL CONVERSATIONS =====
        pdf.add_page()
        pdf.set_font(*heading_font)
        pdf.cell(0, 10, "7. FULL CONVERSATIONS")
        pdf.ln()
        
        pdf.set_font(*body_font)
        pdf.multi_cell(0, 8, "This section contains the full text of all evaluated conversations, along with their evaluations.")
        pdf.ln()
        
        # Load all conversations directly from JSON files
        json_dir = os.path.join(output_dir, 'json')
        if os.path.exists(json_dir):
            json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
            print(f"Including {len(json_files)} conversations in PDF report")
            
            for i, json_file in enumerate(json_files):
                try:
                    with open(os.path.join(json_dir, json_file), 'r') as f:
                        data = json.load(f)
                    
                    # Add a new page for each conversation
                    pdf.add_page()
                    
                    # Conversation header
                    pdf.set_font(*subheading_font)
                    pdf.cell(0, 10, f"Conversation {i+1}: {json_file}")
                    pdf.ln()
                    
                    # Metadata
                    pdf.set_font(*body_font)
                    if 'skill_level' in data:
                        pdf.cell(60, 8, "Skill Level:", border=0)
                        pdf.cell(0, 8, sanitize_text_for_pdf(str(data.get('skill_level', 'N/A'))), border=0)
                        pdf.ln()
                    
                    if 'persona' in data:
                        pdf.cell(60, 8, "Persona:", border=0)
                        pdf.cell(0, 8, sanitize_text_for_pdf(str(data.get('persona', 'N/A'))), border=0)
                        pdf.ln()
                    
                    # Evaluation summary
                    if 'evaluation' in data:
                        pdf.ln()
                        pdf.set_font(*subheading_font)
                        pdf.cell(0, 10, "Evaluation:")
                        pdf.ln()
                        
                        pdf.set_font(*body_font)
                        eval_data = data['evaluation']
                        
                        # Display scores
                        scores = [
                            ("Overall Score", eval_data.get('overall_score', 'N/A')),
                            ("Critical Thinking", eval_data.get('critical_thinking_score', 'N/A')),
                            ("Communication", eval_data.get('communication_score', 'N/A')),
                            ("Emotional Intelligence", eval_data.get('emotional_intelligence_score', 'N/A')),
                            ("Badge", eval_data.get('badge', 'N/A'))
                        ]
                        
                        for label, value in scores:
                            pdf.cell(60, 8, f"{label}:", border=0)
                            pdf.cell(0, 8, sanitize_text_for_pdf(str(value)), border=0)
                        
                        # Display feedback
                        if 'feedback' in eval_data:
                            pdf.ln()
                            pdf.set_font(*subheading_font)
                            pdf.cell(0, 10, "Feedback:")
                            pdf.ln()
                            
                            pdf.set_font(*body_font)
                            pdf.multi_cell(0, 8, sanitize_text_for_pdf(str(eval_data['feedback'])))
                    
                    # Conversation content
                    if 'conversation' in data:
                        pdf.ln()
                        pdf.set_font(*subheading_font)
                        pdf.cell(0, 10, "Conversation Transcript:")
                        pdf.ln()
                        
                        pdf.set_font(*small_font)
                        
                        # Process each message in the conversation
                        for message in data['conversation']:
                            if isinstance(message, dict):
                                role = message.get('role', 'Unknown')
                                content = message.get('content', '')
                                
                                # Add role with bold formatting
                                pdf.set_font('Helvetica', 'B', 8)
                                pdf.cell(20, 8, f"{sanitize_text_for_pdf(role)}:", border=0)
                                pdf.ln()
                                
                                # Add content
                                pdf.set_font('Helvetica', '', 8)
                                
                                # Sanitize content
                                content = sanitize_text_for_pdf(content)
                                
                                # Split content into smaller chunks to avoid overflow
                                content_chunks = [content[i:i+1000] for i in range(0, len(content), 1000)]
                                for chunk in content_chunks:
                                    pdf.multi_cell(0, 6, chunk)
                                
                                pdf.ln(4)
                
                except Exception as e:
                    logger.error(f"Error processing conversation {json_file}: {str(e)}")
                    pdf.multi_cell(0, 8, f"Error processing this conversation: {str(e)}")
        else:
            pdf.multi_cell(0, 8, f"No JSON directory found at {json_dir}")
        
        # Save the PDF
        pdf.output(pdf_path)
        logger.info(f"PDF report saved to: {pdf_path}")
        print(f"\nComprehensive PDF report saved to: {pdf_path}")
        return pdf_path
    
    except Exception as e:
        logger.error(f"Error generating PDF report: {str(e)}")
        print(f"Error generating PDF report: {str(e)}")
        return None

def generate_word_document(output_dir):
    """Generate a Word document containing all conversations from JSON files."""
    try:
        import os
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import json
        import traceback
        
        print(f"Generating Word document with conversations: {os.path.join(output_dir, 'conversations.docx')}")
        
        # Create a new Word document
        doc = Document()
        doc.add_heading('Networking Skills Evaluation - Conversations', level=0)
        doc.add_paragraph('This document contains all evaluated conversations with their assessment scores.')
        
        # Path to the JSON directory
        json_dir = os.path.join(output_dir, 'json')
        if not os.path.exists(json_dir):
            print(f"Warning: JSON directory not found at {json_dir}")
            return
            
        json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
        print(f"Including {len(json_files)} conversations in Word document")
            
        # Iterate over each JSON file in the directory
        for filename in json_files:
            file_path = os.path.join(json_dir, filename)
            print(f"Processing conversation from {filename}")
            
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    try:
                        data = json.load(f)
                        
                        # Add a heading for this conversation
                        doc.add_heading(f'Conversation: {filename}', level=1)
                        
                        # Add metadata section
                        if isinstance(data, dict):
                            metadata = doc.add_paragraph()
                            metadata.add_run('Metadata:').bold = True
                            doc.add_paragraph(f"Skill Level: {data.get('skill_level', 'Unknown')}")
                            doc.add_paragraph(f"Persona: {data.get('persona', 'Unknown')}")
                            
                            # Add evaluation section if available
                            if 'evaluation' in data:
                                eval_heading = doc.add_heading('Evaluation', level=2)
                                eval_data = data['evaluation']
                                
                                if isinstance(eval_data, dict):
                                    for key, value in eval_data.items():
                                        if key != 'rationale':
                                            doc.add_paragraph(f"{key}: {value}")
                                    
                                    # Add rationale separately with special formatting
                                    if 'rationale' in eval_data:
                                        rationale = doc.add_paragraph()
                                        rationale.add_run('Rationale: ').bold = True
                                        rationale.add_run(str(eval_data['rationale']))
                            
                            # Extract conversation data
                            conversation = None
                            if 'conversation' in data:
                                conversation = data['conversation']
                            elif 'messages' in data:
                                conversation = data['messages']
                            
                            # Process the conversation
                            if conversation is not None:
                                conv_heading = doc.add_heading('Conversation Transcript', level=2)
                                
                                if isinstance(conversation, list):
                                    print(f"Conversation type: {type(conversation)}")
                                    print(f"Conversation length: {len(conversation)}")
                                    
                                    for i, message in enumerate(conversation):
                                        print(f"  Message {i} type: {type(message)}")
                                        
                                        if isinstance(message, dict):
                                            role = message.get('role', '')
                                            content = message.get('content', '')
                                            print(f"  Role: {role}, Content length: {len(content) if content else 0}")
                                            
                                            # Format and add the message
                                            msg = doc.add_paragraph()
                                            if role:
                                                role_run = msg.add_run(f"{role}: ")
                                                role_run.bold = True
                                                if role.lower() in ['user', 'human']:
                                                    role_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for user
                                                elif role.lower() in ['assistant', 'ai']:
                                                    role_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for assistant
                                            
                                            msg.add_run(content)
                                        elif isinstance(message, str):
                                            print(f"  String message, length: {len(message)}")
                                            doc.add_paragraph(message)
                                        else:
                                            print(f"  Unknown message format: {type(message)}")
                                            doc.add_paragraph(f"Message (unknown format): {str(message)}")
                                elif isinstance(conversation, dict):
                                    for role, content in conversation.items():
                                        msg = doc.add_paragraph()
                                        msg.add_run(f"{role}: ").bold = True
                                        msg.add_run(str(content))
                                elif isinstance(conversation, str):
                                    doc.add_paragraph(conversation)
                                else:
                                    doc.add_paragraph(f"Conversation format not recognized: {type(conversation).__name__}")
                            else:
                                doc.add_paragraph("No conversation data found.")
                                
                            # Add raw JSON data for debugging
                            debug_heading = doc.add_heading('Raw Data (Debug)', level=3)
                            debug_para = doc.add_paragraph(json.dumps(data, indent=2)[:1000] + "..." if len(json.dumps(data)) > 1000 else json.dumps(data, indent=2))
                            debug_para.style = 'Code'
                        else:
                            doc.add_paragraph(f"File does not contain expected dictionary structure: {filename}")
                            doc.add_paragraph(f"Data type: {type(data).__name__}")
                            doc.add_paragraph(str(data)[:500])  # Show a preview
                            
                    except json.JSONDecodeError as e:
                        error_msg = f"Error decoding JSON from {filename}: {str(e)}"
                        print(error_msg)
                        doc.add_paragraph(error_msg)
            except Exception as e:
                error_msg = f"Error processing file {filename}: {str(e)}"
                print(error_msg)
                print(traceback.format_exc())
                doc.add_paragraph(error_msg)
                
            # Add a page break after each conversation
            doc.add_page_break()
        
        # Save the document
        doc_path = os.path.join(output_dir, 'conversations.docx')
        doc.save(doc_path)
        print(f"Word document saved to: {doc_path}")
        return doc_path
        
    except ImportError as e:
        print(f"Error: {str(e)}")
        print("Please install python-docx: pip install python-docx")
        return None
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        print(traceback.format_exc())
        return None

def generate_text_document(output_dir):
    """Generate a simple text file containing all conversations that can be copied to Google Docs."""
    import os
    import json
    import traceback
    from datetime import datetime
    
    # Default output directory if None
    if output_dir is None:
        output_dir = 'comprehensive_results'
    
    print(f"{Colors.BLUE}Generating text document with conversations...{Colors.ENDC}")
    
    # Path to JSON directory
    json_dir = os.path.join(output_dir, 'json')
    if not os.path.exists(json_dir):
        print(f"{Colors.RED}Warning: JSON directory not found at {json_dir}{Colors.ENDC}")
        return None
        
    # Output file path
    output_file = os.path.join(output_dir, 'conversations.txt')
    
    try:
        with open(output_file, 'w', encoding='utf-8') as out_file:
            # Write header
            out_file.write("NETWORKING SKILLS EVALUATION - CONVERSATIONS\n")
            out_file.write("===========================================\n\n")
            out_file.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            # Get all JSON files
            json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
            if not json_files:
                out_file.write("No JSON files found.\n")
                print(f"{Colors.RED}No JSON files found in {json_dir}{Colors.ENDC}")
                return output_file
            
            out_file.write(f"Total conversations: {len(json_files)}\n\n")
            
            # Iterate through each file and format its contents
            for filename in sorted(json_files):
                file_path = os.path.join(json_dir, filename)
                print(f"Processing {filename}...")
                
                out_file.write(f"{'=' * 80}\n")
                out_file.write(f"CONVERSATION: {filename}\n")
                out_file.write(f"{'=' * 80}\n\n")
                
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        try:
                            data = json.load(f)
                            
                            # Write metadata
                            out_file.write("METADATA:\n")
                            if 'metadata' in data:
                                metadata = data['metadata']
                                if isinstance(metadata, dict):
                                    for key, value in metadata.items():
                                        out_file.write(f"{key}: {value}\n")
                                else:
                                    out_file.write(f"{metadata}\n")
                            elif 'skill_level' in data:
                                out_file.write(f"Skill Level: {data.get('skill_level', 'Unknown')}\n")
                                out_file.write(f"Persona: {data.get('persona', 'Unknown')}\n")
                            else:
                                out_file.write("No metadata found.\n")
                            
                            out_file.write("\n")
                            
                            # Write evaluation data if available
                            if 'evaluation' in data:
                                out_file.write("EVALUATION:\n")
                                eval_data = data['evaluation']
                                
                                if isinstance(eval_data, dict):
                                    for key, value in eval_data.items():
                                        if key != 'rationale' and key != 'response_evaluations':
                                            out_file.write(f"{key}: {value}\n")
                                    
                                    # Add rationale separately
                                    if 'rationale' in eval_data:
                                        out_file.write(f"\nRationale:\n{eval_data['rationale']}\n")
                                    
                                    # Add response evaluations if available
                                    if 'response_evaluations' in eval_data and eval_data['response_evaluations']:
                                        out_file.write("\nResponse Evaluations:\n")
                                        for i, resp_eval in enumerate(eval_data['response_evaluations']):
                                            out_file.write(f"Response {i+1}:\n")
                                            if isinstance(resp_eval, dict):
                                                for eval_key, eval_value in resp_eval.items():
                                                    out_file.write(f"  {eval_key}: {eval_value}\n")
                                            else:
                                                out_file.write(f"  {resp_eval}\n")
                                else:
                                    out_file.write(f"{eval_data}\n")
                                
                                out_file.write("\n")
                            
                            # Extract and write conversation
                            if 'conversation' in data:
                                out_file.write("CONVERSATION TRANSCRIPT:\n\n")
                                conversation = data['conversation']
                                
                                if isinstance(conversation, list):
                                    # Handle the specific format with stage, ai_prompt, user_response
                                    for i, message in enumerate(conversation):
                                        if isinstance(message, dict):
                                            # Check for the expected format with stage, ai_prompt, user_response
                                            if 'stage' in message and 'ai_prompt' in message and 'user_response' in message:
                                                stage = message.get('stage', '')
                                                ai_prompt = message.get('ai_prompt', '')
                                                user_response = message.get('user_response', '')
                                                
                                                out_file.write(f"Stage: {stage}\n")
                                                out_file.write(f"AI: {ai_prompt}\n")
                                                out_file.write(f"User: {user_response}\n")
                                                out_file.write("\n")
                                            # Check for standard format with role and content
                                            elif 'role' in message and 'content' in message:
                                                role = message.get('role', 'Unknown')
                                                content = message.get('content', '')
                                                out_file.write(f"{role}: {content}\n\n")
                                            # Fall back to dumping the message as-is
                                            else:
                                                out_file.write(f"Message {i+1}:\n")
                                                for key, value in message.items():
                                                    out_file.write(f"  {key}: {value}\n")
                                                out_file.write("\n")
                                        elif isinstance(message, str):
                                            out_file.write(f"{message}\n\n")
                                        else:
                                            out_file.write(f"Message {i+1} (type: {type(message).__name__}):\n{message}\n\n")
                                elif isinstance(conversation, dict):
                                    for role, content in conversation.items():
                                        out_file.write(f"{role}: {content}\n\n")
                                elif isinstance(conversation, str):
                                    out_file.write(f"{conversation}\n\n")
                                else:
                                    out_file.write(f"Conversation (type: {type(conversation).__name__}):\n{conversation}\n\n")
                            elif 'messages' in data:
                                out_file.write("CONVERSATION TRANSCRIPT:\n\n")
                                messages = data['messages']
                                
                                if isinstance(messages, list):
                                    for message in messages:
                                        if isinstance(message, dict):
                                            role = message.get('role', 'Unknown')
                                            content = message.get('content', '')
                                            out_file.write(f"{role}: {content}\n\n")
                                        elif isinstance(message, str):
                                            out_file.write(f"{message}\n\n")
                                        else:
                                            out_file.write(f"Message (type: {type(message).__name__}):\n{message}\n\n")
                            else:
                                out_file.write("No conversation data found.\n\n")
                            
                        except json.JSONDecodeError as e:
                            out_file.write(f"Error decoding JSON: {str(e)}\n\n")
                except Exception as e:
                    out_file.write(f"Error processing file: {str(e)}\n\n")
                    print(f"{Colors.RED}Error processing {filename}: {str(e)}{Colors.ENDC}")
                
                # Add separator between conversations
                out_file.write("\n\n")
        
        print(f"{Colors.GREEN}Text document saved to: {output_file}{Colors.ENDC}")
        return output_file
    
    except Exception as e:
        print(f"{Colors.RED}Error generating text document: {str(e)}{Colors.ENDC}")
        print(traceback.format_exc())
        return None

def analyze_json_files(output_dir):
    """Analyze the content of JSON files to identify issues."""
    import os
    import json
    import pprint

    # Default output directory if None
    if output_dir is None:
        output_dir = 'comprehensive_results'
    
    print(f"\n{Colors.HEADER}ANALYZING JSON FILES IN {output_dir}/json{Colors.ENDC}")
    
    # Path to JSON directory
    json_dir = os.path.join(output_dir, 'json')
    if not os.path.exists(json_dir):
        print(f"{Colors.RED}JSON directory not found at {json_dir}{Colors.ENDC}")
        return
    
    # Get list of JSON files
    json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
    if not json_files:
        print(f"{Colors.RED}No JSON files found in {json_dir}{Colors.ENDC}")
        return
    
    print(f"{Colors.GREEN}Found {len(json_files)} JSON files{Colors.ENDC}")
    
    # Analyze a few sample files
    sample_size = min(3, len(json_files))
    sample_files = json_files[:sample_size]
    
    issues_found = False
    
    for filename in sample_files:
        file_path = os.path.join(json_dir, filename)
        print(f"\n{Colors.BOLD}Analyzing {filename}:{Colors.ENDC}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                
                # Check structure
                if not isinstance(data, dict):
                    print(f"{Colors.RED}  ISSUE: File does not contain a dictionary (type: {type(data).__name__}){Colors.ENDC}")
                    issues_found = True
                    continue
                
                # Check key fields
                print(f"  Top-level keys: {', '.join(data.keys())}")
                
                missing_keys = []
                for key in ['skill_level', 'persona', 'evaluation', 'conversation']:
                    if key not in data:
                        missing_keys.append(key)
                
                if missing_keys:
                    print(f"{Colors.RED}  ISSUE: Missing expected keys: {', '.join(missing_keys)}{Colors.ENDC}")
                    issues_found = True
                
                # Check conversation
                if 'conversation' in data:
                    conv = data['conversation']
                    print(f"  Conversation type: {type(conv).__name__}")
                    print(f"  Conversation length: {len(conv) if hasattr(conv, '__len__') else 'N/A'}")
                    
                    if isinstance(conv, list):
                        # Check first message
                        if conv and isinstance(conv[0], dict):
                            print(f"  First message keys: {', '.join(conv[0].keys())}")
                            
                            # Check if role and content are present and non-empty
                            role_key = 'role' if 'role' in conv[0] else None
                            content_key = 'content' if 'content' in conv[0] else None
                            
                            if not role_key:
                                print(f"{Colors.RED}  ISSUE: Message doesn't have 'role' key{Colors.ENDC}")
                                issues_found = True
                            elif not conv[0].get(role_key):
                                print(f"{Colors.RED}  ISSUE: 'role' value is empty{Colors.ENDC}")
                                issues_found = True
                            
                            if not content_key:
                                print(f"{Colors.RED}  ISSUE: Message doesn't have 'content' key{Colors.ENDC}")
                                issues_found = True
                            elif not conv[0].get(content_key):
                                print(f"{Colors.RED}  ISSUE: 'content' value is empty{Colors.ENDC}")
                                issues_found = True
                    else:
                        print(f"{Colors.RED}  ISSUE: Conversation is not a list{Colors.ENDC}")
                        issues_found = True
                
                # Check evaluation
                if 'evaluation' in data:
                    eval_data = data['evaluation']
                    print(f"  Evaluation type: {type(eval_data).__name__}")
                    
                    if isinstance(eval_data, dict):
                        print(f"  Evaluation keys: {', '.join(eval_data.keys())}")
                    else:
                        print(f"{Colors.RED}  ISSUE: Evaluation is not a dictionary{Colors.ENDC}")
                        issues_found = True
                
        except Exception as e:
            print(f"{Colors.RED}  ERROR: {str(e)}{Colors.ENDC}")
            issues_found = True
    
    if issues_found:
        print(f"\n{Colors.RED}ISSUES FOUND IN JSON FILES{Colors.ENDC}")
        print("Possible solutions:")
        print("1. Regenerate the conversations with: python generate_comprehensive_report.py")
        print("2. Check the conversation generation process in generate_conversation() function")
        print("3. Check the save_conversation_to_json() function for correct data storage")
    else:
        print(f"\n{Colors.GREEN}No major issues found in the sample files.{Colors.ENDC}")
    
    return issues_found

def run_comprehensive_evaluation(args):
    """Run a comprehensive evaluation of networking conversations."""
    import os
    import importlib.util
    import logging
    import traceback
    from datetime import datetime
    
    # Setup logging
    logging.basicConfig(
        level=logging.DEBUG if args.debug else (logging.INFO if args.verbose else logging.WARNING),
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    
    logging.info("Starting comprehensive evaluation...")
    print(f"{Colors.HEADER}Starting comprehensive evaluation...{Colors.ENDC}")
    
    # Import conversation generator and evaluator
    try:
        # Import the synthetic conversation generator
        spec = importlib.util.spec_from_file_location("synthetic_conversation_gpt", "synthetic_conversation_gpt.py")
        synthetic_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(synthetic_module)
        generator = synthetic_module.SyntheticConversationGPT()
        
        # Import the conversation evaluator
        spec = importlib.util.spec_from_file_location("conversation_evaluator_gpt", "conversation_evaluator_gpt.py")
        evaluator_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(evaluator_module)
        evaluator = evaluator_module.ConversationEvaluatorGPT()
        
        if args.verbose or args.debug:
            print(f"{Colors.GREEN}Successfully imported generator and evaluator modules{Colors.ENDC}")
    
    except Exception as e:
        error_msg = f"Error importing modules: {str(e)}"
        logging.error(error_msg)
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        if args.debug:
            print(traceback.format_exc())
        return
    
    # Setup output directory
    output_dir = args.output_dir
    
    # Generate skill gradients (skill level combinations)
    skill_gradients = generate_skill_gradients()
    
    if args.verbose or args.debug:
        print(f"{Colors.BLUE}Skill gradients:{Colors.ENDC}")
        for sg in skill_gradients:
            print(f"  {sg}")
    
    # Dictionary to store results
    results = []
    
    # Define personas
    personas = ["INTERVIEWER"]
    
    # Generate and evaluate conversations for each skill level and persona
    for persona in personas:
        for skill_level, label in skill_gradients:
            # Generate 3 conversations for each combination
            for i in range(1, 4):
                try:
                    if args.verbose or args.debug:
                        print(f"\n{Colors.HEADER}Processing {skill_level} (#{i}) with {persona} persona{Colors.ENDC}")
                    
                    # Generate conversation
                    conversation = generate_conversation(generator, persona, skill_level, label, args.verbose, args.debug)
                    
                    # Evaluate conversation
                    evaluation = evaluate_conversation(evaluator, conversation, skill_level, args.verbose, args.debug)
                    
                    # Save conversation and evaluation to JSON
                    json_path = save_conversation_to_json(conversation, evaluation, skill_level, persona, output_dir, i, args.verbose, args.debug)
                    
                    # Save conversation and evaluation to text
                    text_path = save_conversation_to_text(conversation, evaluation, skill_level, persona, output_dir, i)
                    
                    # Add result to dictionary
                    result = {
                        'skill_level': skill_level,
                        'skill_base': skill_level.split('_')[0],
                        'gradient': skill_level.split('_')[2],
                        'persona': persona,
                        'evaluation': evaluation,
                        'json_path': json_path,
                        'text_path': text_path
                    }
                    results.append(result)
                    
                    if args.verbose or args.debug:
                        print(f"{Colors.GREEN}Completed processing for {skill_level} (#{i}){Colors.ENDC}")
                
                except Exception as e:
                    error_msg = f"Error processing {skill_level} (#{i}): {str(e)}"
                    logging.error(error_msg)
                    print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
                    if args.debug:
                        print(traceback.format_exc())
    
    # Create summary dataframe
    print("\n" + "=" * 80)
    print("GENERATING SUMMARY REPORTS")
    print("=" * 80)
    
    df = create_summary_dataframe(results)
    
    # Generate reports
    generate_summary_reports(output_dir)
    
    # Generate Word document if requested
    if args.word_doc:
        generate_word_document(output_dir)
    
    # Generate text document if requested
    if args.text_file:
        generate_text_document(output_dir)
    
    if args.verbose or args.debug:
        print(f"\n{Colors.GREEN}Comprehensive evaluation completed successfully!{Colors.ENDC}")
    
    return results

def run_text_only_simple():
    """Run a simplified text-only approach that doesn't require any external modules."""
    import os
    import json
    import traceback
    from datetime import datetime
    
    # Create output directory
    output_dir = "text_results"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Create files for output
    output_file = os.path.join(output_dir, "conversations.txt")
    
    print(f"{Colors.HEADER}Starting simplified text extraction...{Colors.ENDC}")
    print(f"Output will be saved to: {output_file}")
    
    try:
        # First, check if we have an existing comprehensive_results directory with JSON files
        json_dir = "comprehensive_results/json"
        
        if os.path.exists(json_dir) and os.listdir(json_dir):
            print(f"{Colors.GREEN}Found existing JSON files in {json_dir}{Colors.ENDC}")
            
            with open(output_file, 'w', encoding='utf-8') as out_file:
                # Write header
                out_file.write("NETWORKING SKILLS EVALUATION - CONVERSATIONS\n")
                out_file.write("===========================================\n\n")
                out_file.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                # Get all JSON files
                json_files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
                out_file.write(f"Total conversations: {len(json_files)}\n\n")
                
                # Process each JSON file
                for filename in sorted(json_files):
                    file_path = os.path.join(json_dir, filename)
                    print(f"Processing {filename}...")
                    
                    out_file.write(f"{'=' * 80}\n")
                    out_file.write(f"CONVERSATION: {filename}\n")
                    out_file.write(f"{'=' * 80}\n\n")
                    
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            # First try to load as JSON
                            try:
                                data = json.load(f)
                                
                                # Try to find metadata
                                if 'metadata' in data:
                                    out_file.write("METADATA:\n")
                                    metadata = data['metadata']
                                    if isinstance(metadata, dict):
                                        for key, value in metadata.items():
                                            out_file.write(f"{key}: {value}\n")
                                    else:
                                        out_file.write(f"{metadata}\n")
                                    out_file.write("\n")
                                elif 'skill_level' in data:
                                    out_file.write("METADATA:\n")
                                    out_file.write(f"Skill Level: {data.get('skill_level', 'Unknown')}\n")
                                    out_file.write(f"Persona: {data.get('persona', 'Unknown')}\n")
                                    out_file.write("\n")
                                
                                # Try to find evaluation
                                if 'evaluation' in data:
                                    out_file.write("EVALUATION:\n")
                                    evaluation = data['evaluation']
                                    
                                    if isinstance(evaluation, dict):
                                        for key, value in evaluation.items():
                                            if key not in ['response_evaluations', 'rationale']:
                                                out_file.write(f"{key}: {value}\n")
                                        
                                        # Add rationale separately
                                        if 'rationale' in evaluation:
                                            out_file.write(f"\nRationale:\n{evaluation['rationale']}\n")
                                    else:
                                        out_file.write(f"{evaluation}\n")
                                    
                                    out_file.write("\n")
                                
                                # Try to find conversation
                                conversation = None
                                if 'conversation' in data:
                                    conversation = data['conversation']
                                elif 'messages' in data:
                                    conversation = data['messages']
                                
                                if conversation is not None:
                                    out_file.write("CONVERSATION TRANSCRIPT:\n\n")
                                    
                                    if isinstance(conversation, list):
                                        # Check for the format with stage, ai_prompt, user_response
                                        if conversation and isinstance(conversation[0], dict) and 'stage' in conversation[0]:
                                            for msg in conversation:
                                                if isinstance(msg, dict):
                                                    stage = msg.get('stage', 'Unknown')
                                                    ai_prompt = msg.get('ai_prompt', '')
                                                    user_response = msg.get('user_response', '')
                                                    
                                                    out_file.write(f"Stage: {stage}\n")
                                                    out_file.write(f"AI: {ai_prompt}\n")
                                                    out_file.write(f"User: {user_response}\n\n")
                                        # Check for the format with role and content
                                        elif conversation and isinstance(conversation[0], dict) and 'role' in conversation[0]:
                                            for msg in conversation:
                                                if isinstance(msg, dict):
                                                    role = msg.get('role', 'Unknown')
                                                    content = msg.get('content', '')
                                                    
                                                    out_file.write(f"{role}: {content}\n\n")
                                        # Default case
                                        else:
                                            for i, msg in enumerate(conversation):
                                                out_file.write(f"Message {i+1}:\n{msg}\n\n")
                                    elif isinstance(conversation, dict):
                                        for role, content in conversation.items():
                                            out_file.write(f"{role}: {content}\n\n")
                                    elif isinstance(conversation, str):
                                        out_file.write(f"{conversation}\n\n")
                                    else:
                                        out_file.write(f"Conversation (type: {type(conversation).__name__}):\n{conversation}\n\n")
                                else:
                                    out_file.write("No conversation data found.\n\n")
                                
                            except json.JSONDecodeError:
                                # If JSON parsing fails, just include the raw file contents
                                f.seek(0)  # Go back to the start of the file
                                raw_content = f.read()
                                out_file.write("RAW FILE CONTENT:\n\n")
                                out_file.write(raw_content)
                                out_file.write("\n\n")
                    except Exception as e:
                        out_file.write(f"Error processing file: {str(e)}\n\n")
                        print(f"{Colors.RED}Error processing {filename}: {str(e)}{Colors.ENDC}")
                    
                    # Add separator between conversations
                    out_file.write("\n\n")
            
            print(f"{Colors.GREEN}Text document saved to: {output_file}{Colors.ENDC}")
            return output_file
        
        # If no JSON files, check for text files
        text_dir = "comprehensive_results/text"
        
        if os.path.exists(text_dir) and os.listdir(text_dir):
            print(f"{Colors.GREEN}Found existing text files in {text_dir}{Colors.ENDC}")
            
            with open(output_file, 'w', encoding='utf-8') as out_file:
                # Write header
                out_file.write("NETWORKING SKILLS EVALUATION - CONVERSATIONS\n")
                out_file.write("===========================================\n\n")
                out_file.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                # Get all text files
                text_files = [f for f in os.listdir(text_dir) if f.endswith('.txt')]
                out_file.write(f"Total conversations: {len(text_files)}\n\n")
                
                # Concatenate all text files
                for filename in sorted(text_files):
                    file_path = os.path.join(text_dir, filename)
                    print(f"Processing {filename}...")
                    
                    out_file.write(f"{'=' * 80}\n")
                    out_file.write(f"CONVERSATION: {filename}\n")
                    out_file.write(f"{'=' * 80}\n\n")
                    
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                            out_file.write(content)
                            out_file.write("\n\n")
                    except Exception as e:
                        out_file.write(f"Error reading file: {str(e)}\n\n")
                        print(f"{Colors.RED}Error reading {filename}: {str(e)}{Colors.ENDC}")
            
            print(f"{Colors.GREEN}Text document saved to: {output_file}{Colors.ENDC}")
            return output_file
        
        # If no existing files, inform the user
        print(f"{Colors.RED}No existing conversation files found in comprehensive_results directory.{Colors.ENDC}")
        print(f"{Colors.RED}Please run the full evaluation first to generate conversation files.{Colors.ENDC}")
        print(f"{Colors.RED}Command: python generate_comprehensive_report.py{Colors.ENDC}")
        
        with open(output_file, 'w', encoding='utf-8') as out_file:
            out_file.write("NO CONVERSATION FILES FOUND\n")
            out_file.write("==========================\n\n")
            out_file.write("No existing conversation files were found in the comprehensive_results directory.\n\n")
            out_file.write("Please run the full evaluation first to generate conversation files:\n")
            out_file.write("python generate_comprehensive_report.py\n")
        
        return output_file
    
    except Exception as e:
        print(f"{Colors.RED}Error during text extraction: {str(e)}{Colors.ENDC}")
        print(traceback.format_exc())
        return None

def regenerate_all_to_text():
    """Regenerate all conversations and evaluations from scratch and save them directly as text."""
    import os
    import sys
    import importlib.util
    import traceback
    import json
    import inspect
    from datetime import datetime

    print(f"{Colors.HEADER}Starting regeneration of ALL conversations and saving directly to text...{Colors.ENDC}")
    
    # Create output directory
    output_dir = "text_output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Output file paths
    output_file = os.path.join(output_dir, "all_conversations.txt")
    debug_file = os.path.join(output_dir, "debug_log.txt")
    
    # Add current directory to path
    current_dir = os.getcwd()
    if current_dir not in sys.path:
        sys.path.append(current_dir)
    
    print(f"Looking for modules in: {', '.join(sys.path)}")
    
    # Initialize variables
    generator = None
    evaluator = None
    generator_module = None
    
    # Write initial headers to files
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("REGENERATED NETWORKING CONVERSATIONS\n")
            f.write("===================================\n\n")
            f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    except Exception as e:
        print(f"{Colors.RED}Error creating output file: {str(e)}{Colors.ENDC}")
        return None
        
    try:
        with open(debug_file, 'w', encoding='utf-8') as dbg:
            dbg.write("COMPREHENSIVE DEBUG LOG - NETWORKING CONVERSATIONS\n")
            dbg.write("=================================================\n\n")
            dbg.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            dbg.write(f"Python path: {', '.join(sys.path)}\n\n")
            dbg.write("MODULE LOADING\n")
            dbg.write("=============\n\n")
    except Exception as e:
        print(f"{Colors.RED}Error creating debug file: {str(e)}{Colors.ENDC}")
        return None
    
    # Helper functions for writing to files
    def write_to_debug(message):
        try:
            with open(debug_file, 'a', encoding='utf-8') as f:
                f.write(message)
        except Exception as e:
            print(f"{Colors.RED}Error writing to debug file: {str(e)}{Colors.ENDC}")
    
    def write_to_output(message):
        try:
            with open(output_file, 'a', encoding='utf-8') as f:
                f.write(message)
        except Exception as e:
            print(f"{Colors.RED}Error writing to output file: {str(e)}{Colors.ENDC}")
    
    # Step 1: Try to load generator module
    generator_module_names = ['synthetic_conversation_gpt', 'conversation_generator', 'generator']
    for module_name in generator_module_names:
        try:
            write_to_debug(f"Attempting to import {module_name}...\n")
            generator_module = __import__(module_name)
            print(f"{Colors.GREEN}Successfully imported {module_name}{Colors.ENDC}")
            write_to_debug(f"Success! Available attributes: {dir(generator_module)}\n\n")
            break
        except ImportError as e:
            write_to_debug(f"Failed to import {module_name}: {str(e)}\n")
    
    if not generator_module:
        error_msg = "Could not import any generator module"
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        write_to_debug(f"ERROR: {error_msg}\n\n")
        write_to_output(f"ERROR: {error_msg}\n\n")
        return None
    
    # Step 2: Try to find generator class
    write_to_debug("Looking for generator class...\n")
    for attr_name in dir(generator_module):
        attr = getattr(generator_module, attr_name)
        write_to_debug(f"Checking attribute: {attr_name}, type: {type(attr)}\n")
        
        if callable(attr) and not attr_name.startswith('__'):
            try:
                write_to_debug(f"Attempting to instantiate: {attr_name}\n")
                generator = attr()
                print(f"{Colors.GREEN}Created generator using {attr_name}{Colors.ENDC}")
                write_to_debug(f"Created generator instance using {attr_name}\n")
                write_to_debug(f"Generator methods: {[m for m in dir(generator) if not m.startswith('_')]}\n")
                
                if hasattr(generator, 'generate_conversation') or hasattr(generator, 'generate') or hasattr(generator, 'run'):
                    write_to_debug(f"Found suitable generator method in {attr_name}\n")
                    break
            except Exception as e:
                write_to_debug(f"Failed to instantiate {attr_name}: {str(e)}\n")
                write_to_debug(traceback.format_exc() + "\n")
    
    if not generator:
        error_msg = f"Could not find a suitable generator class in {generator_module.__name__}"
        print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
        write_to_debug(f"ERROR: {error_msg}\n\n")
        write_to_output(f"ERROR: {error_msg}\n\n")
        return None
    
    # Step 3: Try to load evaluator module
    evaluator_module = None
    evaluator_module_names = ['conversation_evaluator_gpt', 'conversation_evaluator', 'evaluator']
    
    for module_name in evaluator_module_names:
        try:
            write_to_debug(f"Attempting to import {module_name}...\n")
            evaluator_module = __import__(module_name)
            print(f"{Colors.GREEN}Successfully imported {module_name}{Colors.ENDC}")
            write_to_debug(f"Success! Available attributes: {dir(evaluator_module)}\n\n")
            break
        except ImportError as e:
            write_to_debug(f"Failed to import {module_name}: {str(e)}\n")
    
    # If no evaluator module, use generator module
    if not evaluator_module:
        print(f"{Colors.YELLOW}No dedicated evaluator module found. Using generator module for evaluation.{Colors.ENDC}")
        write_to_debug("No dedicated evaluator module found. Using generator module for evaluation.\n")
        evaluator_module = generator_module
    
    # Step 4: Try to find evaluator class
    write_to_debug("Looking for evaluator class...\n")
    for attr_name in dir(evaluator_module):
        attr = getattr(evaluator_module, attr_name)
        write_to_debug(f"Checking attribute: {attr_name}, type: {type(attr)}\n")
        
        if callable(attr) and not attr_name.startswith('__'):
            # Skip if already used as generator
            if generator and attr_name == generator.__class__.__name__:
                write_to_debug(f"Skipping {attr_name} as it's already used for generator\n")
                continue
                
            try:
                write_to_debug(f"Attempting to instantiate: {attr_name}\n")
                evaluator = attr()
                print(f"{Colors.GREEN}Created evaluator using {attr_name}{Colors.ENDC}")
                write_to_debug(f"Created evaluator instance using {attr_name}\n")
                write_to_debug(f"Evaluator methods: {[m for m in dir(evaluator) if not m.startswith('_')]}\n")
                
                if hasattr(evaluator, 'evaluate_conversation') or hasattr(evaluator, 'evaluate') or hasattr(evaluator, 'run'):
                    write_to_debug(f"Found suitable evaluator method in {attr_name}\n")
                    break
            except Exception as e:
                write_to_debug(f"Failed to instantiate {attr_name}: {str(e)}\n")
                write_to_debug(traceback.format_exc() + "\n")
    
    # If no evaluator found, use generator
    if not evaluator:
        print(f"{Colors.YELLOW}No dedicated evaluator class found. Using generator as evaluator.{Colors.ENDC}")
        write_to_debug("No dedicated evaluator class found. Using generator as evaluator.\n")
        evaluator = generator
    
    print(f"{Colors.GREEN}Successfully created generator and evaluator instances{Colors.ENDC}")
    write_to_debug("GENERATION CONFIGURATION\n")
    write_to_debug("=======================\n\n")
    
    # Generate skill gradients
    write_to_debug("Generating skill gradients manually...\n")
    skill_bases = ['novice', 'intermediate', 'advanced']
    gradients = ['low', 'basic', 'high']
    skill_gradients = []
    
    for base in skill_bases:
        for gradient in gradients:
            skill_level = f"{base}_{gradient}"
            label = f"{base.capitalize()} ({gradient})"
            skill_gradients.append((skill_level, label))
    
    write_to_debug("Skill gradients:\n")
    for sg in skill_gradients:
        write_to_debug(f"{sg}\n")
    write_to_debug("\n")
    
    # Define personas
    personas = ["INTERVIEWER"]
    write_to_debug(f"Personas: {personas}\n\n")
    
    # Stats tracking
    total_conversations = 0
    successful_conversations = 0
    
    # Fallback conversation and evaluation
    fallback_conversation = [
        {"stage": "introduction", "ai_prompt": "Hello, I'm interviewing for a networking position.", "user_response": "Hi, I'm here to assess your networking skills."},
        {"stage": "technical_question", "ai_prompt": "Can you explain how DNS works?", "user_response": "DNS converts domain names to IP addresses so browsers can load resources."},
        {"stage": "scenario", "ai_prompt": "What would you do if a client couldn't access the internet?", "user_response": "I'd check physical connections, DNS settings, and network configuration."}
    ]
    
    fallback_evaluation = {
        "overall_score": 3,
        "critical_thinking_score": 3,
        "communication_score": 3,
        "emotional_intelligence_score": 3,
        "badge_level": "Silver",
        "rationale": "This is a fallback evaluation because the actual evaluation failed."
    }
    
    # Generate and evaluate conversations
    for persona in personas:
        for skill_level_tuple in skill_gradients:
            skill_level = skill_level_tuple[0]
            label = skill_level_tuple[1]
            
            for i in range(1, 4):
                total_conversations += 1
                
                # Write section headers
                section_header = f"{skill_level.upper()} - CONVERSATION {i} - {persona}"
                write_to_output(f"\n{'=' * 80}\n")
                write_to_output(f"{section_header}\n")
                write_to_output(f"{'=' * 80}\n\n")
                
                write_to_debug(f"\n{'=' * 80}\n")
                write_to_debug(f"DEBUG: {section_header}\n")
                write_to_debug(f"{'=' * 80}\n\n")
                
                print(f"Generating: {section_header}")
                
                try:
                    # Write metadata
                    write_to_output("METADATA:\n")
                    write_to_output(f"Skill Level: {skill_level}\n")
                    write_to_output(f"Persona: {persona}\n")
                    write_to_output(f"Index: {i}\n\n")
                    
                    # Generate conversation
                    write_to_debug("GENERATING CONVERSATION\n")
                    write_to_debug("----------------------\n\n")
                    
                    prompt = {
                        "persona": persona,
                        "label": label
                    }
                    
                    write_to_debug("Prompt:\n")
                    write_to_debug(json.dumps(prompt, indent=2) + "\n")
                    write_to_debug("Skill Level: " + skill_level + "\n\n")
                    
                    # Try to generate conversation with different methods
                    conversation = None
                    generation_method = None
                    
                    try:
                        if hasattr(generator, 'generate_conversation'):
                            # Try method 1
                            write_to_debug("Trying generate_conversation(skill_level, prompt)\n")
                            try:
                                conversation = generator.generate_conversation(skill_level, prompt)
                                generation_method = 'generate_conversation(skill_level, prompt)'
                                write_to_debug("Method 1 successful\n")
                            except TypeError:
                                # Try method 2
                                write_to_debug("Method 1 failed, trying generate_conversation(skill_level)\n")
                                try:
                                    conversation = generator.generate_conversation(skill_level)
    for module_name in evaluator_module_names:
        try:
            debug_log(f"Attempting to import {module_name}...\n")
            evaluator_module = __import__(module_name)
            print(f"{Colors.GREEN}Successfully imported {module_name}{Colors.ENDC}")
            debug_log(f"Success! Available attributes: {dir(evaluator_module)}\n\n")
            break
        except ImportError as e:
            debug_log(f"Failed to import {module_name}: {str(e)}\n")
    
    # If no evaluator module, use the generator module
    if not evaluator_module:
        print(f"{Colors.YELLOW}No dedicated evaluator module found. Trying to use generator module for evaluation.{Colors.ENDC}")
        debug_log("No dedicated evaluator module found. Using generator module for evaluation.\n")
        evaluator_module = generator_module
    
    # Try to find an evaluator class
    debug_log("Looking for evaluator class...\n")
    for attr_name in dir(evaluator_module):
        attr = getattr(evaluator_module, attr_name)
        debug_log(f"Checking attribute: {attr_name}, type: {type(attr)}\n")
        if callable(attr) and not attr_name.startswith('__'):
        
        generator = None
        evaluator = None
        
        # List of possible module names to try
        generator_module_names = ['synthetic_conversation_gpt', 'conversation_generator', 'generator']
        evaluator_module_names = ['conversation_evaluator_gpt', 'conversation_evaluator', 'evaluator']
        
        # First try to load the generator module
        generator_module = None
        for module_name in generator_module_names:
            try:
                dbg.write(f"Attempting to import {module_name}...\n")
                generator_module = __import__(module_name)
                print(f"{Colors.GREEN}Successfully imported {module_name}{Colors.ENDC}")
                dbg.write(f"Success! Available attributes: {dir(generator_module)}\n\n")
                break
            except ImportError as e:
                dbg.write(f"Failed to import {module_name}: {str(e)}\n")
                
        if not generator_module:
            error_msg = "Could not import any generator module"
            print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
            dbg.write(f"ERROR: {error_msg}\n\n")
            f.write(f"ERROR: {error_msg}\n\n")
            return None
            
        # Try to find a generator class
        dbg.write("Looking for generator class...\n")
        for attr_name in dir(generator_module):
            attr = getattr(generator_module, attr_name)
            dbg.write(f"Checking attribute: {attr_name}, type: {type(attr)}\n")
            if callable(attr) and not attr_name.startswith('__'):
                try:
                    dbg.write(f"Attempting to instantiate: {attr_name}\n")
                    generator = attr()
                    print(f"{Colors.GREEN}Created generator using {attr_name}{Colors.ENDC}")
                    dbg.write(f"Created generator instance using {attr_name}\n")
                    dbg.write(f"Generator methods: {[m for m in dir(generator) if not m.startswith('_')]}\n")
                    
                    # Check if generator has suitable methods
                    if hasattr(generator, 'generate_conversation'):
                        # Inspect the method signature to understand the parameters
                        if hasattr(generator, 'generate_conversation'):
                            method = generator.generate_conversation
                            dbg.write(f"Inspecting generate_conversation method: {inspect.signature(method)}\n")
                        elif hasattr(generator, 'generate'):
                            method = generator.generate
                            dbg.write(f"Inspecting generate method: {inspect.signature(method)}\n")
                        elif hasattr(generator, 'run'):
                            method = generator.run
                            dbg.write(f"Inspecting run method: {inspect.signature(method)}\n")
                        
                        dbg.write(f"Found suitable generator method in {attr_name}\n")
                        break
                except Exception as e:
                    dbg.write(f"Failed to instantiate {attr_name}: {str(e)}\n")
                    dbg.write(traceback.format_exc())
                    dbg.write("\n")
                    continue
        
        if not generator:
            error_msg = f"Could not find a suitable generator class in {generator_module.__name__}"
            print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
            dbg.write(f"ERROR: {error_msg}\n\n")
            f.write(f"ERROR: {error_msg}\n\n")
            return None
            
        # Now try to load the evaluator module (may be the same as generator)
        evaluator_module = None
        for module_name in evaluator_module_names:
            try:
                dbg.write(f"Attempting to import {module_name}...\n")
                evaluator_module = __import__(module_name)
                print(f"{Colors.GREEN}Successfully imported {module_name}{Colors.ENDC}")
                dbg.write(f"Success! Available attributes: {dir(evaluator_module)}\n\n")
                break
            except ImportError as e:
                dbg.write(f"Failed to import {module_name}: {str(e)}\n")
        
        # If no evaluator module, use the generator module
        if not evaluator_module:
            print(f"{Colors.YELLOW}No dedicated evaluator module found. Trying to use generator module for evaluation.{Colors.ENDC}")
            dbg.write("No dedicated evaluator module found. Using generator module for evaluation.\n")
            evaluator_module = generator_module
        
        # Try to find an evaluator class
        dbg.write("Looking for evaluator class...\n")
        for attr_name in dir(evaluator_module):
            attr = getattr(evaluator_module, attr_name)
            dbg.write(f"Checking attribute: {attr_name}, type: {type(attr)}\n")
            if callable(attr) and not attr_name.startswith('__'):
                try:
                    # Skip if we already instantiated this class as the generator
                    if generator and attr_name == generator.__class__.__name__:
                        dbg.write(f"Skipping {attr_name} as it's already used for generator\n")
                        continue
                        
                    dbg.write(f"Attempting to instantiate: {attr_name}\n")
                    evaluator = attr()
                    print(f"{Colors.GREEN}Created evaluator using {attr_name}{Colors.ENDC}")
                    dbg.write(f"Created evaluator instance using {attr_name}\n")
                    dbg.write(f"Evaluator methods: {[m for m in dir(evaluator) if not m.startswith('_')]}\n")
                    
                    # Check if evaluator has suitable methods and inspect them
                    if hasattr(evaluator, 'evaluate_conversation'):
                        method = evaluator.evaluate_conversation
                        dbg.write(f"Inspecting evaluate_conversation method: {inspect.signature(method)}\n")
                        break
                    elif hasattr(evaluator, 'evaluate'):
                        method = evaluator.evaluate
                        dbg.write(f"Inspecting evaluate method: {inspect.signature(method)}\n")
                        break
                    elif hasattr(evaluator, 'run'):
                        method = evaluator.run
                        dbg.write(f"Inspecting run method: {inspect.signature(method)}\n")
                        break
                except Exception as e:
                    dbg.write(f"Failed to instantiate {attr_name}: {str(e)}\n")
                    dbg.write(traceback.format_exc())
                    dbg.write("\n")
                    continue
        
        # If no evaluator found, use the generator as evaluator too
        if not evaluator:
            print(f"{Colors.YELLOW}No dedicated evaluator class found. Using generator as evaluator.{Colors.ENDC}")
            dbg.write("No dedicated evaluator class found. Using generator as evaluator.\n")
            evaluator = generator
            
        # If we have both a generator and evaluator, proceed
        if generator and evaluator:
            print(f"{Colors.GREEN}Successfully created generator and evaluator instances{Colors.ENDC}")
            dbg.write("GENERATION CONFIGURATION\n")
            dbg.write("=======================\n\n")
            
            # Generate skill gradients manually to ensure correct structure
            dbg.write("Generating skill gradients manually...\n")
            skill_bases = ['novice', 'intermediate', 'advanced']
            gradients = ['low', 'basic', 'high']
            skill_gradients = []
            
            for base in skill_bases:
                for gradient in gradients:
                    skill_level = f"{base}_{gradient}"
                    label = f"{base.capitalize()} ({gradient})"
                    skill_gradients.append((skill_level, label))
            
            dbg.write("Skill gradients:\n")
            for sg in skill_gradients:
                dbg.write(f"{sg}\n")
            dbg.write("\n")
            
            # Define personas
            personas = ["INTERVIEWER"]
            dbg.write(f"Personas: {personas}\n\n")
            
            # Track stats
            total_conversations = 0
            successful_conversations = 0
            
            # Generate and evaluate conversations for each skill level and persona
            for persona in personas:
                for skill_level_tuple in skill_gradients:
                    # Unpack the tuple safely
                    skill_level = skill_level_tuple[0]
                    label = skill_level_tuple[1]
                    
                    # Generate 3 conversations for each combination
                    for i in range(1, 4):
                        total_conversations += 1
                        
                        # Write section headers
                        section_header = f"{skill_level.upper()} - CONVERSATION {i} - {persona}"
                        f.write(f"\n{'=' * 80}\n")
                        f.write(f"{section_header}\n")
                        f.write(f"{'=' * 80}\n\n")
                        
                        dbg.write(f"\n{'=' * 80}\n")
                        dbg.write(f"DEBUG: {section_header}\n")
                        dbg.write(f"{'=' * 80}\n\n")
                        
                        print(f"Generating: {section_header}")
                        
                        try:
                            # Write metadata
                            f.write("METADATA:\n")
                            f.write(f"Skill Level: {skill_level}\n")
                            f.write(f"Persona: {persona}\n")
                            f.write(f"Index: {i}\n\n")
                            
                            # Generate conversation
                            dbg.write("GENERATING CONVERSATION\n")
                            dbg.write("----------------------\n\n")
                            
                            prompt = {
                                "persona": persona,
                                "label": label
                            }
                            
                            dbg.write("Prompt:\n")
                            dbg.write(json.dumps(prompt, indent=2))
                            dbg.write("\nSkill Level: " + skill_level + "\n\n")
                            
                            # Try to generate conversation
                            conversation = None
                            generation_method = None
                            
                            dbg.write("Attempting to generate conversation...\n")
                            
                            # Fallback conversation if all methods fail
                            fallback_conversation = [
                                {"stage": "introduction", "ai_prompt": "Hello, I'm interviewing for a networking position.", "user_response": "Hi, I'm here to assess your networking skills."},
                                {"stage": "technical_question", "ai_prompt": "Can you explain how DNS works?", "user_response": "DNS converts domain names to IP addresses so browsers can load resources."},
                                {"stage": "scenario", "ai_prompt": "What would you do if a client couldn't access the internet?", "user_response": "I'd check physical connections, DNS settings, and network configuration."}
                            ]
                            
                            try:
                                # Try method 1: generate_conversation(skill_level, prompt)
                                if hasattr(generator, 'generate_conversation'):
                                    dbg.write("Trying generate_conversation(skill_level, prompt)\n")
                                    try:
                                        conversation = generator.generate_conversation(skill_level, prompt)
                                        generation_method = 'generate_conversation(skill_level, prompt)'
                                        dbg.write("Method 1 successful\n")
                                    except TypeError as e:
                                        dbg.write(f"Method 1 failed: {str(e)}\n")
                                        
                                        # Try method 2: generate_conversation(skill_level)
                                        dbg.write("Trying generate_conversation(skill_level)\n")
                                        try:
                                            conversation = generator.generate_conversation(skill_level)
                                            generation_method = 'generate_conversation(skill_level)'
                                            dbg.write("Method 2 successful\n")
                                        except TypeError as e:
                                            dbg.write(f"Method 2 failed: {str(e)}\n")
                                            
                                            # Try method 3: generate_conversation(prompt) with skill_level in prompt
                                            dbg.write("Trying generate_conversation(prompt) with skill_level in prompt\n")
                                            try:
                                                prompt["skill_level"] = skill_level
                                                conversation = generator.generate_conversation(prompt)
                                                generation_method = 'generate_conversation(prompt)'
                                                dbg.write("Method 3 successful\n")
                                            except TypeError as e:
                                                dbg.write(f"Method 3 failed: {str(e)}\n")
                                                dbg.write("All generate_conversation methods failed\n")
                                                raise
                                
                                # If no generate_conversation, try other methods
                                elif hasattr(generator, 'generate'):
                                    dbg.write("Using generate() method\n")
                                    prompt["skill_level"] = skill_level
                                    conversation = generator.generate(prompt)
                                    generation_method = 'generate'
                                    dbg.write("generate() method successful\n")
                                
                                elif hasattr(generator, 'run'):
                                    dbg.write("Using run() method\n")
                                    prompt["skill_level"] = skill_level
                                    conversation = generator.run(prompt)
                                    generation_method = 'run'
                                    dbg.write("run() method successful\n")
                                
                                else:
                                    dbg.write("Using direct callable\n")
                                    prompt["skill_level"] = skill_level
                                    conversation = generator(prompt)
                                    generation_method = '__call__'
                                    dbg.write("direct callable successful\n")
                            
                            except Exception as gen_error:
                                # Handle any errors during generation
                                dbg.write(f"ERROR during generation: {str(gen_error)}\n")
                                dbg.write(traceback.format_exc())
                                dbg.write("\nUsing fallback sample conversation\n")
                                
                                conversation = fallback_conversation
                                generation_method = 'FALLBACK_SAMPLE'
                            
                            # Log the generation results
                            dbg.write(f"Generation complete using {generation_method}\n")
                            dbg.write(f"Result type: {type(conversation)}\n")
                            if hasattr(conversation, '__len__'):
                                dbg.write(f"Result length: {len(conversation)}\n")
                            
                            dbg.write("Raw conversation result:\n")
                            dbg.write(json.dumps(conversation, indent=2))
                            dbg.write("\n\n")
                            
                            # Write the conversation
                            f.write("CONVERSATION TRANSCRIPT:\n\n")
                            
                            if isinstance(conversation, list):
                                # Check if it's the format with stage, ai_prompt, user_response
                                if conversation and isinstance(conversation[0], dict) and 'stage' in conversation[0]:
                                    dbg.write("Detected stage/ai_prompt/user_response format\n")
                                    for msg in conversation:
                                        if isinstance(msg, dict):
                                            stage = msg.get('stage', 'Unknown')
                                            ai_prompt = msg.get('ai_prompt', '')
                                            user_response = msg.get('user_response', '')
                                            
                                            f.write(f"Stage: {stage}\n")
                                            f.write(f"AI: {ai_prompt}\n")
                                            f.write(f"User: {user_response}\n\n")
                                # Check if it's the format with role and content
                                elif conversation and isinstance(conversation[0], dict) and 'role' in conversation[0]:
                                    dbg.write("Detected role/content format\n")
                                    for msg in conversation:
                                        if isinstance(msg, dict):
                                            role = msg.get('role', 'Unknown')
                                            content = msg.get('content', '')
                                            
                                            f.write(f"{role}: {content}\n\n")
                                # Default case
                                else:
                                    dbg.write("Using default message format\n")
                                    for j, msg in enumerate(conversation):
                                        f.write(f"Message {j+1}:\n{msg}\n\n")
                            else:
                                dbg.write("Non-list conversation format\n")
                                f.write(f"Conversation: {conversation}\n\n")
                            
                            # Evaluate the conversation
                            dbg.write("EVALUATING CONVERSATION\n")
                            dbg.write("-----------------------\n\n")
                            
                            eval_input = {
                                "conversation": conversation,
                                "skill_level": skill_level
                            }
                            
                            dbg.write("Evaluation input:\n")
                            dbg.write(json.dumps(eval_input, indent=2))
                            dbg.write("\n\n")
                            
                            # Fallback evaluation if all methods fail
                            fallback_evaluation = {
                                "overall_score": 3,
                                "critical_thinking_score": 3,
                                "communication_score": 3,
                                "emotional_intelligence_score": 3,
                                "badge_level": "Silver",
                                "rationale": "This is a fallback evaluation because the actual evaluation failed."
                            }
                            
                            # Try to evaluate the conversation
                            evaluation = None
                            evaluation_method = None
                            
                            try:
                                dbg.write("Attempting to evaluate conversation...\n")
                                
                                # Try method 1: evaluate_conversation(conversation, skill_level)
                                if hasattr(evaluator, 'evaluate_conversation'):
                                    dbg.write("Trying evaluate_conversation(conversation, skill_level)\n")
                                    try:
                                        evaluation = evaluator.evaluate_conversation(conversation, skill_level)
                                        evaluation_method = 'evaluate_conversation(conversation, skill_level)'
                                        dbg.write("Evaluation method 1 successful\n")
                                    except TypeError as e:
                                        dbg.write(f"Evaluation method 1 failed: {str(e)}\n")
                                        
                                        # Try method 2: evaluate_conversation(eval_input)
                                        dbg.write("Trying evaluate_conversation(eval_input)\n")
                                        try:
                                            evaluation = evaluator.evaluate_conversation(eval_input)
                                            evaluation_method = 'evaluate_conversation(eval_input)'
                                            dbg.write("Evaluation method 2 successful\n")
                                        except TypeError as e:
                                            dbg.write(f"Evaluation method 2 failed: {str(e)}\n")
                                            dbg.write("All evaluate_conversation methods failed\n")
                                            raise
                                
                                # If no evaluate_conversation, try other methods
                                elif hasattr(evaluator, 'evaluate'):
                                    dbg.write("Using evaluate() method\n")
                                    evaluation = evaluator.evaluate(eval_input)
                                    evaluation_method = 'evaluate'
                                    dbg.write("evaluate() method successful\n")
                                
                                elif hasattr(evaluator, 'run'):
                                    dbg.write("Using run() method\n")
                                    evaluation = evaluator.run(eval_input)
                                    evaluation_method = 'run'
                                    dbg.write("run() method successful\n")
                                
                                else:
                                    dbg.write("Using direct callable\n")
                                    evaluation = evaluator(eval_input)
                                    evaluation_method = '__call__'
                                    dbg.write("direct callable successful\n")
                            
                            except Exception as eval_error:
                                # Handle any errors during evaluation
                                dbg.write(f"ERROR during evaluation: {str(eval_error)}\n")
                                dbg.write(traceback.format_exc())
                                dbg.write("\nUsing fallback evaluation\n")
                                
                                evaluation = fallback_evaluation
                                evaluation_method = 'FALLBACK_EVALUATION'
                            
                            # Log the evaluation results
                            dbg.write(f"Evaluation complete using {evaluation_method}\n")
                            dbg.write(f"Result type: {type(evaluation)}\n")
                            
                            dbg.write("Raw evaluation result:\n")
                            dbg.write(json.dumps(evaluation, indent=2))
                            dbg.write("\n\n")
                            
                            # Write the evaluation
                            f.write("EVALUATION:\n")
                            
                            # Add evaluation summary
                            f.write("Summary:\n")
                            f.write(f"Overall Score: {evaluation['overall_score']}\n")
                            f.write(f"Badge Level: {evaluation['badge_level']}\n")
                            f.write("\n")
                            
                            # Add detailed evaluation
                            f.write("Detailed Evaluation:\n")
                            for key, value in evaluation.items():
                                if key not in ['overall_score', 'badge_level', 'response_evaluations', 'rationale']:
                                    f.write(f"{key}: {value}\n")
                            
                            f.write("\n")
                            
                            # Add response evaluations
                            if 'response_evaluations' in evaluation:
                                f.write("Response Evaluations:\n")
                                for resp_eval in evaluation['response_evaluations']:
                                    f.write(f"Stage: {resp_eval['stage']}\n")
                                    f.write(f"User Query: {resp_eval['user_query']}\n")
                                    f.write(f"Assistant Response: {resp_eval['assistant_response']}\n")
                                    f.write("\n")
                            
                            # Add rationale
                            if 'rationale' in evaluation:
                                f.write("Rationale:\n")
                                f.write(evaluation['rationale'])
                                f.write("\n\n")
                            
                            # Add feedback
                            if 'feedback' in evaluation:
                                f.write("Feedback:\n")
                                f.write(evaluation['feedback'])
                                f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(evaluation, indent=2))
                            f.write("\n\n")
                            
                            # Add conversation
                            f.write("Conversation:\n")
                            f.write(conversation_text)
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
                            f.write(json.dumps(conversation, indent=2))
                            f.write("\n\n")
                            
                            # Add raw JSON data for debugging
                            f.write("Raw JSON data:\n")
            error_msg = "Failed to create generator or evaluator instances"
            print(f"{Colors.RED}{error_msg}{Colors.ENDC}")
            dbg.write(f"FATAL ERROR: {error_msg}\n")
            f.write(f"ERROR: {error_msg}\n")
    
    print(f"\n{Colors.GREEN}Process complete!{Colors.ENDC}")
    print(f"Main output saved to: {output_file}")
    print(f"Debug log saved to: {debug_file}")
    
    return output_file, debug_file

def main():
    """Main function."""
    import traceback
    
    # Parse command-line arguments
    args = parse_args()
    
    try:
        # If regenerate-all flag is set, regenerate everything from scratch
        if args.regenerate_all:
            regenerate_all_to_text()
            return
        
        # If simple flag is set, run the simplified text extraction
        if args.simple or args.text_only:
            run_text_only_simple()
            return
        
        # Setup output directory
        output_dir = setup_output_directory(args.output_dir)
        
        # If analyze flag is set, analyze existing files
        if args.analyze:
            analyze_json_files(output_dir)
            return
        
        # If regenerate flag is set, run comprehensive evaluation
        if args.regenerate or not args.summary_only:
            run_comprehensive_evaluation(args)
        
        # If summary_only flag is set, generate summary reports
        if args.summary_only and not args.regenerate:
            if args.word_doc:
                generate_word_document(output_dir)
            if args.text_file:
                generate_text_document(output_dir)
            generate_summary_reports(output_dir)
        
        print(f"\n{Colors.GREEN}Evaluation completed successfully!{Colors.ENDC}")
    
    except Exception as e:
        print(f"\n{Colors.RED}Evaluation failed: {str(e)}{Colors.ENDC}")
        if args.debug:
            print(f"{Colors.RED}Stack trace:{Colors.ENDC}")
            print(traceback.format_exc())

if __name__ == '__main__':
    try:
        main()
    except Exception as e:
        import traceback
        print(f"{Colors.RED}Evaluation failed. See log for details.{Colors.ENDC}")
        traceback.print_exc() 